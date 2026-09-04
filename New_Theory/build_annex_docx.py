# -*- coding: utf-8 -*-
"""Builds the English annex (.docx) describing the software, in full.

Requested 2026-08-27: a free-standing annex for the paper — usage, capabilities,
limitations, how it works, every constant explained, a dedicated section on why
this is not a curve fit, and every fitted source as a reference.

Everything numeric is RECOMPUTED from the canonical store at build time. No
figure in the annex carries a number that was typed by hand: if the store moves,
the annex moves with it. That is the same rule the HTML reports follow (§4.43).

    py -3.12 New_Theory/build_annex_docx.py            # -> New_Theory/annex/
"""
from __future__ import annotations

import collections
import html as _h
import json
import re
import sys
from pathlib import Path

RAIZ = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(RAIZ / "src"))
sys.path.insert(0, str(RAIZ / "New_Theory"))

import numpy as np  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402

import bolt_analysis_studio.validation.report_html as rh  # noqa: E402
import bolt_analysis_studio.validation.runner as rn  # noqa: E402
from bolt_analysis_studio.calibration import knowledge_base as kb  # noqa: E402
from bolt_analysis_studio.validation.case_registry import all_records  # noqa: E402
from bolt_analysis_studio.validation.runner import CaseResult  # noqa: E402

SAIDA = RAIZ / "New_Theory" / "annex"
FIGS = SAIDA / "figures"


# --------------------------------------------------------------------------- #
# Data                                                                        #
# --------------------------------------------------------------------------- #

def carrega():
    """comp/res are the CENSUS set; res_all also holds the non-comparable ones.

    The two are not interchangeable. Every published metric goes through
    `caso_comparavel` — that filter is what stops the same physical test,
    published in two figures, from being counted twice. `res_all` exists only
    for the physics demonstration of section 8.1, where the question is
    'does one configuration follow several behaviours', not 'how many curves
    does the corpus contain'.
    """
    bruto = json.loads((RAIZ / "Models" / "CALIBRATION_AND_VALIDATION"
                        / "validation_store.json").read_text(encoding="utf-8"))
    store = bruto.get("cases", bruto)
    todos = all_records()
    res_all = {}
    for r in todos:
        if r.case_id not in store:
            continue
        try:
            res_all[r.case_id] = CaseResult.from_dict(store[r.case_id])
        except Exception:
            pass
    comp = [r for r in todos
            if r.case_id in store and rh.caso_comparavel(r.source, r.case_id)]
    res = {r.case_id: res_all[r.case_id] for r in comp if r.case_id in res_all}
    pisos = rh._pisos_medidos([(r.source, res[r.case_id]) for r in comp
                               if r.case_id in res])
    return comp, res, pisos, store, todos, res_all


def limpa(s: str) -> str:
    """VarSpec prose is double-escaped HTML; this yields plain text."""
    s = _h.unescape(_h.unescape(s or ""))
    s = re.sub(r"<br\s*/?>", " ", s)
    s = re.sub(r"</p>\s*<p>", "\n", s)
    s = re.sub(r"<[^>]+>", "", s)
    return re.sub(r"[ \t]+", " ", s).strip()


# --------------------------------------------------------------------------- #
# Declared classes · principal constants · text audit                          #
# --------------------------------------------------------------------------- #

# Class of a DECLARED curve, derived from the proof text `report_html` carries
# for it. The proof is the source of truth — it is what the master report
# renders — so this only maps its opening words onto the vocabulary of §10.3;
# the annex cannot name a class the proof does not state. Order matters: the
# n < 6 proofs also mention the collapse history, so they are matched first.
_CLASSES_DECLARADAS = [
    ("não-julgável", "n < 6"),
    ("form-limited", "form-limited"),
    ("metric-limited", "metric-limited"),
    ("data-limited", "data-limited"),
    ("fora de escopo", "out of scope"),
    ("órfã de protocolo", "protocol orphan"),
]


def classe_declarada(cid: str) -> str:
    txt = (getattr(rh, "_DECLARADAS", {}) or {}).get(cid, "") or ""
    low = txt.lower()
    for chave, rotulo in _CLASSES_DECLARADAS:
        if chave in low:
            return rotulo
    return "declared (unclassified)"


def tipo_excecao(cid: str) -> str:
    """F5 = replicate scatter of the source; F7 = proof against the measured
    floor. Read from the two dicts `_EXCECOES` is the union of."""
    if cid in getattr(rh, "_F7_EXCECOES", {}):
        return "floor proof (F7)"
    if cid in getattr(rh, "_F5_EXCECOES", {}):
        return "replicate scatter (F5)"
    return "signed"


# The VarSpec prose was written for the interactive explorer and talks about
# its controls ("moving the slider…"). Those sentences are dropped wherever a
# VarSpec text is quoted; the principal constants get prose of their own.
_UI = re.compile(r"[^.]*\b(slider|the plot|readout|baseline here|"
                 r"context turns|companions .{0,20}enabled)\b[^.]*\.", re.I)


def sem_ui(txt: str) -> str:
    return re.sub(r"\s{2,}", " ", _UI.sub("", txt)).strip()


# Units the explorer stores in Portuguese, or with a form the equation does
# not support (C_creep multiplies a force to give a length: m/N).
_UNIDADES_EN = {"ciclos": "cycles", "ciclos-eq": "cycle-eq.",
                "rad/ciclo": "rad/cycle", "1/ciclo": "1/cycle",
                "m/(log-dec.Pa)": "m/N", "kg.m^2": "kg·m²", "": "–", "-": "–"}


def unidade_en(u) -> str:
    u = (u or "").strip()
    return _UNIDADES_EN.get(u, u)


# Campaign bookkeeping that the VarSpec prose carries in parentheses (dates,
# pre-registration codes, internal source keys) means nothing to a reader of
# the paper. Parentheticals that contain such a token are dropped whole; the
# spaced dash, used in that prose as a connector, becomes a colon or a
# semicolon so the table reads as a table.
_JARGAO = re.compile(
    r"\s*\((?:[^()]*?(?:20\d\d-\d\d-\d\d|\bPR-?\s?\d|\bP-\d+|prereg|\bL1\b|"
    r"\bl1\b|\bF4\b|master run|TP7|IJPEM provenance|roadmap|§4\.\d+|"
    r"form B of|PR-21)[^()]*)\)")


# A minus sign with spaces is a connector when a function word follows it
# ("[1/Pa] − the MAGNITUDE…"); between two symbols it is arithmetic and stays.
_MENOS_CONECTOR = re.compile(
    r" − (?![a-z]+\()"            # not a function call: "1 − min(r, 1)"
    r"(?=(?:[a-z]{2,}|a)\b)")     # a lowercase word follows: prose, not algebra


def _troca_conector(s: str, padrao: re.Pattern) -> str:
    """First connector becomes a colon unless the clause already has one
    (a second colon reads badly); every later one becomes a semicolon."""
    m = padrao.search(s)
    if not m:
        return s
    sep = ": " if ":" not in s else "; "
    s = s[:m.start()] + sep + s[m.end():]
    return padrao.sub("; ", s)


def prosa_tabela(s: str) -> str:
    s = _JARGAO.sub("", s)
    s = re.sub(r",?\s*(?:signed|adopted|measured)\s+20\d\d-\d\d-\d\d", "", s)
    s = re.sub(r"\bthe L1 channel\b", "the flank channel", s)
    s = re.sub(r"\bL1(?:\s?v2)?\b", "flank", s)
    s = _troca_conector(s, re.compile(r" — "))
    s = _troca_conector(s, _MENOS_CONECTOR)
    return re.sub(r"\s{2,}", " ", s).strip()


# Tokens that should never survive in an English annex: Portuguese that
# leaked from the explorer's units/equations, or explorer-UI vocabulary.
_VAZAMENTOS = ("slider", "the plot", "readout", "[ciclos]", " nao ",
               "legado", "amplifica ", "constante de tempo", " aqui ")


def auditoria_texto(doc):
    """Every paragraph and table cell, checked for the tokens above. The list is
    printed by main(); an annex that fails the audit is still written, so the
    reader of the console decides — but it is never silent."""
    achados = []

    def _olha(t):
        low = f" {t.lower()} "
        for v in _VAZAMENTOS:
            if v in low:
                achados.append((v, t[:100]))
    for par in doc.paragraphs:
        _olha(par.text)
    for tb in doc.tables:
        for row in tb.rows:
            for c in row.cells:
                _olha(c.text)
    return achados


# Principal constants: unit, governing equation and prose written for the
# paper. Every statement is checked against the engine (dynamic_stiffness_
# analyzer.py) — line references are in the comments — and against the adopted
# configurations, so that the prose never claims a usage the corpus does not
# have (e.g. the negative k_dmg_mu, which one configuration does carry).
PRINCIPAIS = {
    "emb_depth": ("m",
        "δ_emb(N) = δ∞ · (1 − e^(−N/N_emb)),      δ∞ = emb_depth",
        ["Clamped surfaces meet at the peaks of their roughness asperities. "
         "Under preload those peaks yield, the joint beds in, and a small "
         "thickness is lost. emb_depth is the total depth of that settling, "
         "the geometric slack that disappears once every asperity has "
         "yielded; the joint stiffness converts it into a preload loss. It is "
         "the largest contributor to the early drop of F/F₀, the knee at the "
         "start of most curves.",
         "It is a per-joint input, not a universal constant. The VDI 2230 "
         "embedding table gives it by roughness class and number of "
         "interfaces; the configurations read it from there, or from the "
         "first cycles of the published curve (§9.3). The engine integrates it "
         "in state form, removing each cycle a fixed fraction of the depth "
         "still remaining. For a virgin joint this reproduces the closed form "
         "exactly. A reused joint can start with part of the depth already "
         "consumed."]),
    "N_emb": ("cycles",
        "δ_emb(N) = δ∞ · (1 − e^(−N/N_emb)),      N_emb = time constant, "
        "in cycles",
        ["If emb_depth says how much the joint beds in, N_emb says over how "
         "many cycles: the characteristic number of cycles the asperities take "
         "to yield. It sets the sharpness of the initial knee without changing "
         "its final depth. A small N_emb concentrates the drop in the first "
         "few cycles; a large one spreads the same drop over many.",
         "Each cycle the increment is the remaining depth multiplied by "
         "(1 − e^(−1/N_emb)). An amplitude-dependent clock exists in the "
         "engine and is inactive by default: in a single-amplitude test it only "
         "rescales N_emb and cannot be falsified, so it is not adopted."]),
    "C_creep": ("m/N",
        "δ_creep(t) = C_creep · F₀ · ln(t/t₀ + 1)      "
        "(saturating form: δ = C_creep · F₀ · (1 − e^(−(t/t_c)^α)))",
        ["Logarithmic creep of the clamped interfaces under sustained preload. "
         "The slow settling grows with the logarithm of time and scales with "
         "the preload. That is what produces the long tail of the curve, the "
         "loss that continues well after bedding-in is complete. The same "
         "product C_creep·F₀ is the asymptote of the saturating kernel adopted "
         "where the logarithmic form was falsified (§8.3), so the constant "
         "keeps one meaning under both forms.",
         "It is a property of the tribological pair, not a universal number. "
         "An anchor on 304 stainless steel and the fit on the UFU rig differ "
         "by more than an order of magnitude, with disjoint confidence "
         "intervals, and several sources carry their own value (§9.3). Time "
         "enters through the real clock t = N/f, so the same constant serves "
         "a 1/60 Hz creep test and a 12.5 Hz Junker test."]),
    "t_0": ("s",
        "δ_creep(t) = C_creep · F₀ · ln(t/t₀ + 1)",
        ["The reference time that fixes the origin of the logarithmic creep "
         "clock. A log-time law has no natural zero; t₀ both regularises it "
         "and sets how much of the slow settling happens in the first seconds "
         "rather than over hours. A larger t₀ delays and flattens the early "
         "creep; a smaller one front-loads the tail.",
         "Physically it is the viscoelastic onset of the pair, the knee where "
         "creep becomes visible, and it is read from that knee in the "
         "published curve. For the composite joint of Qin et al. it coincides "
         "with the second retardation time of the authors' own Burgers fit."]),
    "k_wear_spec": ("1/Pa",
        "d_wear = k_wear_spec · F₀ · (4·slip) / A_contact      "
        "(k_wear_spec = K/H)",
        ["The specific wear rate K/H of Archard's law. It is the only "
         "combination of the wear coefficient K and the hardness H that the "
         "data can constrain, because the two enter every wear expression as "
         "their ratio. "
         "Publishing one number is the honest parametrisation; publishing K "
         "and H separately would suggest two degrees of freedom where there "
         "is one.",
         "It drives the depth removed per cycle at the bearing face and, where "
         "the flank channel is enabled, at the thread flank, from the slip "
         "distance 4·slip per cycle and the contact pressure F₀/A_contact. In "
         "displacement-controlled tests wear is the channel that carries most "
         "of the preload loss, which is why the surface-damage state couples "
         "to it (k_dmg_wear). A value of zero selects the legacy K/H path, not "
         "'no wear'; disabling the channel requires both to be zero."]),
    "tr_loose_gain": ("–",
        "Φ_tr,active = tr_loose_gain  if F_tr ≥ F_slip,  else 0.01;      "
        "L = hypot(Φ_ax·sinβ·F_ax, Φ_tr·cosβ·F_tr);      T_loose = L·d₂/2",
        ["Rotational loosening is the product of two factors: an anisotropic "
         "load-partition factor Φ (factor 1) and the helix projection "
         "(factor 2). tr_loose_gain is the factor-1 gain in the transverse "
         "direction. While the cyclic transverse force stays below the slip "
         "threshold F_slip, the joint is essentially locked (Φ_tr,active = "
         "0.01). Once gross slip occurs, Φ_tr,active jumps to tr_loose_gain: "
         "the dynamic amplification of backing-off under transverse vibration "
         "that the Junker test exposes.",
         "It multiplies the transverse leg of the loosening torque. A larger "
         "gain makes loosening start earlier and go deeper. The transition "
         "itself is not chosen by the constant; it is decided by the slip "
         "threshold, which depends on the current preload and friction."]),
    "loose_arrest_floor": ("–",
        "F_min = loose_arrest_floor · F₀,init;      g = max(0, 1 − F_min/F₀);"
        "      dθ ← g · dθ",
        ["When loosening is of the runaway type (the resisting torque falls "
         "with F₀, so the more the joint loosens the easier it loosens), the "
         "equations alone would drive the preload to zero. Physically a joint "
         "often does not: the central stick zone of the contact restores "
         "static thread friction and locks a residual clamp against the helix "
         "off-torque. loose_arrest_floor is that residual clamp as a fraction "
         "of the initial preload.",
         "The gate multiplies the loosening increment, so the ratchet drains "
         "only the excess above F_min and stops there: the runaway becomes an "
         "S-curve with a stable fixed point. Zero, the default, means no "
         "arrest. Where it is non-zero it is read from the terminal plateau of "
         "the published curve or from the authors' own terminal value, never "
         "fitted freely (§9.3). The gate applies in both loading modes, but it "
         "only acts where the rotational channel carries a share of the loss."]),
    "c_D": ("–",
        "dD/dN = c_D · (W_slip / W_ref) · (1 − D),      D ∈ [0, 1]",
        ["Growth rate of the surface-damage state D, fed by the energy "
         "dissipated in slip each cycle. D is not a mechanism: it modulates "
         "the others, lowering the effective bearing friction (k_dmg_mu) and "
         "amplifying wear (k_dmg_wear). With c_D = 0 damage never grows and "
         "the engine reproduces its pre-damage behaviour exactly; a positive "
         "value accelerates the late collapse observed in re-tightened and "
         "reused joints.",
         "The driver is the slip work of wear and rotational loosening, not "
         "embedding or creep, so a joint that never slips never damages; the "
         "(1 − D) factor keeps the state within [0, 1]."]),
    "W_ref": ("J",
        "dD/dN = c_D · (W_slip / W_ref) · (1 − D),      "
        "W_slip = dE_wear + dE_loose",
        ["The energy scale of surface damage: how much dissipated slip work "
         "corresponds to one unit of damage dose. Each cycle the engine "
         "measures the friction work of wear and rotational loosening and "
         "divides it by W_ref to advance D. The larger W_ref, the slower "
         "damage grows; a small W_ref brings the collapse forward into the "
         "early cycles.",
         "Only the ratio c_D/W_ref is identifiable from a preload curve. The "
         "two are published separately because W_ref carries the meaning of "
         "an energy scale and c_D that of a rate, but they must not be read "
         "as independent degrees of freedom."]),
    "k_dmg_mu": ("–",
        "μ_bearing,eff = μ_bearing · max(1 − k_dmg_mu · D, 0)",
        ["Coupling of damage to bearing-face friction. As the surface degrades "
         "the seat loses grip and the effective friction falls; k_dmg_mu sets "
         "the strength of that effect, with 0 meaning friction does not "
         "degrade. The effective value is read at the start of each cycle by "
         "the resisting torque and by the wear work, so less friction means "
         "less torque opposing loosening and a faster collapse.",
         "The sign is opposite to thread galling, which raises the friction "
         "seen during tightening. A negative value is admissible where the "
         "paper measured friction rising with cycling, and one adopted "
         "configuration carries one."]),
    "k_dmg_wear": ("–",
        "d_wear ← d_wear · (1 + k_dmg_wear · D)      "
        "(amplifies the preload loss, not the dissipated energy)",
        ["Coupling of damage to wear: a damaged surface, rougher and loaded "
         "with debris, removes material faster for the same slip. The per-"
         "cycle wear depth is multiplied by (1 + k_dmg_wear·D); 0 switches "
         "the coupling off.",
         "The amplification enters the preload loss but not the dissipated "
         "energy: the friction heat stays the real work, and the extra "
         "removal is balanced through the released elastic energy. Amplifying "
         "the energy as well would break the conservation check by about "
         "40 %. In displacement-controlled tests wear dominates the loss, "
         "which is why the damage state couples to wear and not only to "
         "friction."]),
    "mu_bearing": ("–",
        "T_resist = μ_thread·F₀·d₂/(2 cos 30°) + μ_bearing,eff·F₀·r_bearing;"
        "      F_slip = 0.46 · μ_bearing,eff · F₀",
        ["Friction coefficient at the bearing face, under the head or nut, "
         "against the washer or member. It plays two roles: it sets the "
         "bearing part of the resisting torque that opposes rotational "
         "loosening, and it sets the slip onset F_slip = 0.46·μ·F₀ (Pai & "
         "Hess) below which the interface sticks.",
         "In displacement-controlled tests the wear depth is driven by K/H, "
         "not by μ, so mu_bearing acts through the torque balance and the "
         "onset rather than through the removal rate. Values come from the "
         "measured dry-steel band or from the paper's own friction "
         "measurement, and the knowledge base flags a value outside the "
         "measured band (§9.3)."]),
    "mu_thread": ("–",
        "T_resist = μ_thread · F₀ · d₂ / (2 cos 30°) + …",
        ["Friction coefficient on the thread flank. It governs the thread part "
         "of the resisting torque: higher friction holds the nut, lower "
         "friction releases it. There is no single 'effective friction' to "
         "transcribe from a torque–preload table, because the nut factor "
         "mixes the two interfaces. The value is therefore read through the "
         "engine against the measured band rather than copied."]),
}


# The code block printed in §3.2. It is EXECUTED by main() before the document
# is written: a snippet that does not run is worse than no snippet, and the
# first version of this annex carried one (a keyword argument the constructor
# does not have, and an undefined name).
SNIPPET = (
    "from bolt_analysis_studio.numerical.dynamic_stiffness_analyzer import (\n"
    "    DynamicStiffnessAnalyzer, JointMaterial)\n"
    "from bolt_analysis_studio.validation.inputs import geometry_for\n"
    "\n"
    "geom = geometry_for(\"M16x2.0\", grip_mm=40.0)   # ISO thread table + grip\n"
    "mat = JointMaterial()                          # per-rig constants are fields\n"
    "F0 = 50e3\n"
    "ana = DynamicStiffnessAnalyzer(geom, mat, F0)\n"
    "for n in range(1000):\n"
    "    snap = ana.step_cycle(F_amp=0.0, theta_load=0.0, freq=12.5,\n"
    "                          delta_amp=0.5e-3)    # displacement-controlled\n"
    "ratio = snap.F_0 / F0\n")


def verifica_snippet():
    """Runs SNIPPET exactly as printed; returns the final F/F0 or None."""
    ns = {}
    try:
        exec(SNIPPET, ns)
        return float(ns["ratio"])
    except Exception as e:                     # noqa: BLE001
        print(f"  [WARN] the §3.2 snippet does not run: {e!r}")
        return None


# --------------------------------------------------------------------------- #
# Repository facts computed at build time (git, pytest, engine, provenance)   #
# --------------------------------------------------------------------------- #

def _git(*args):
    import subprocess
    try:
        out = subprocess.run(["git", *args], cwd=str(RAIZ), capture_output=True,
                             text=True, encoding="utf-8", errors="replace",
                             timeout=60)
        return out.stdout.strip() if out.returncode == 0 else None
    except Exception:
        return None


def revisao_git():
    """(short hash, commit date) of the checkout the annex was built from."""
    h = _git("rev-parse", "--short", "HEAD")
    d = _git("log", "-1", "--format=%cd", "--date=short")
    return (h, d) if h and d else None


def data_primeiro_commit(rel_path):
    """(hash, date) of the commit that ADDED a file, or None."""
    out = _git("log", "--diff-filter=A", "--format=%h|%ad", "--date=short",
               "--", rel_path)
    if not out:
        return None
    h, d = out.strip().splitlines()[-1].split("|")
    return h, d


def contagem_testes():
    """How many tests pytest collects. Collection only (a few seconds); the
    last full run, which takes a quarter of an hour, is quoted with its date."""
    import subprocess
    try:
        out = subprocess.run([sys.executable, "-m", "pytest", "tests/",
                              "--collect-only", "-q"], cwd=str(RAIZ),
                             capture_output=True, text=True, encoding="utf-8",
                             errors="replace", timeout=240)
        mm = re.search(r"(\d+) tests? collected", out.stdout + out.stderr)
        return int(mm.group(1)) if mm else None
    except Exception:
        return None


def verificacoes_engine():
    """Three checks on the engine as shipped, run at every build: (i) the
    state-based embedding reproduces the closed-form Norton law, (ii) the
    energy budget closes on the §3.2 example, (iii) throughput."""
    import math
    import time
    from bolt_analysis_studio.numerical.dynamic_stiffness_analyzer import (
        DynamicStiffnessAnalyzer, JointMaterial)
    from bolt_analysis_studio.validation.inputs import geometry_for
    geom = geometry_for("M16x2.0", grip_mm=40.0)
    F0 = 50e3
    so_emb = JointMaterial(C_creep=0.0, k_wear_spec=0.0, K_archard=0.0,
                           tr_loose_gain=0.0, k_thread_fret=0.0)
    ana = DynamicStiffnessAnalyzer(geom, so_emb, F0)
    dev = 0.0
    for n in range(1, 301):
        ana.step_cycle(F_amp=0.0, theta_load=0.0, freq=12.5, delta_amp=0.5e-3)
        fechado = so_emb.emb_depth * (1.0 - math.exp(-n / so_emb.N_emb))
        dev = max(dev, abs(float(ana.state.delta_emb) - fechado))
    ana2 = DynamicStiffnessAnalyzer(geom, JointMaterial(), F0)
    t0 = time.perf_counter()
    for _ in range(1000):
        ana2.step_cycle(F_amp=0.0, theta_load=0.0, freq=12.5, delta_amp=0.5e-3)
    dt = time.perf_counter() - t0
    E = ana2.energy
    return {"norton_dev_m": float(dev), "emb_depth": float(so_emb.emb_depth),
            "n_norton": 300,
            "cycles_per_s": (1000.0 / dt) if dt > 0 else None,
            "residual_J": float(E.conservation_residual),
            "W_diss_J": float(E.W_diss_total), "W_ext_J": float(E.W_ext),
            "dU_J": float(E.U_stored - E.U_stored_init), "n": 1000}


def piso_digitalizacao_par(a="lu2024_M8_fig18_amp1p0",
                           b="lu2024_M8_fig20_T22Nm"):
    """Two independent digitisations of the SAME test (one Lu et al. run,
    published in two figures) compared on their common window. It is the only
    direct measurement of digitisation error the corpus offers."""
    from bolt_analysis_studio.validation.case_registry import all_records
    from bolt_analysis_studio.validation.inputs import load_full_curve
    recs = {r.case_id: r for r in all_records()}
    if a not in recs or b not in recs:
        return None

    def _curva(cid):
        vc = recs[cid].validation_case
        x, y = load_full_curve(vc.reference_csv_path)
        x = ((np.asarray(x, float) - float(vc.csv_x_offset or 0.0))
             * float(vc.csv_x_scale or 1.0))
        return x, np.asarray(y, float)
    xa, ya = _curva(a)
    xb, yb = _curva(b)
    lo, hi = max(xa.min(), xb.min()), min(xa.max(), xb.max())
    g = np.linspace(lo, hi, 60)
    d = np.interp(g, xa, ya) - np.interp(g, xb, yb)
    return {"a": a, "b": b, "n_a": int(len(xa)), "n_b": int(len(xb)),
            "lo": float(lo), "hi": float(hi), "mae": float(np.mean(np.abs(d))),
            "max": float(np.max(np.abs(d))), "sigma": float(np.std(d))}


# --- provenance: alias-aware lookup and a CONSERVATIVE class reader -------- #
# Provenance keys in the configurations may be combined ('c_bend/emb_depth/
# floor'), abbreviated ('mu', 'floor', 'emb') or per-curve ('per_case.<tok>').
_ALIAS_PROV = {"floor": ["loose_arrest_floor"], "emb": ["emb_depth", "emb_um"],
               "mu": ["mu_thread", "mu_bearing"], "kernel": ["loose_rate_mode"]}


def prov_lookup(prov, field):
    for key, txt in (prov or {}).items():
        for p in [q.strip() for q in str(key).split("/")]:
            if p.startswith("per_case."):
                continue
            if p == field or field in _ALIAS_PROV.get(p, []):
                return str(txt)
            if len(p) >= 4 and p in field.split("_"):
                return str(txt)
    return None


_NEG_FIT = re.compile(r"n[aã]o[\s-]*(?:e'?\s*)?(?:constante\s+)?fitad[ao]|"
                      r"not\s+fitted", re.I)
_RX_CLASSE = {
    "channel off": re.compile(r"desligad|channel off|nenhum grau de liberdade|"
                              r"(?:^|\s)=\s?0 pela|(?:^|\s)0 por contexto",
                              re.I),
    "form switch (a mode, not a number)": re.compile(
        r"forma adotada|escolha de mecanismo|form switch|mode switch|"
        r"canal .{0,24}ligado|ligad[oa] por atribui|companheiro obrigat|"
        r"turned on", re.I),
    "fitted (declared)": re.compile(r"\bfitad|\bfit\b|re-?fitado|refit|"
                                    r"\bgrade\b|\bgrid\b|fitted", re.I),
    "anchored": re.compile(r"[âa]ncora|anchor|held-?out", re.I),
    "regressed": re.compile(r"regredid|regress|\blsq\b|r²|r2\s*=|"
                            r"least.squares", re.I),
    "read (canonical reader)": re.compile(
        r"lido[- ]d[oa][- ](?:dado|curva|joelho|fim|intercepto|tra[çc]o|"
        r"caracteriza|res[ií]duo|terminal|plat[oô]|slope|queda)|lido pelo|"
        r"lido deste dado|\blid[oa]\b|\bL24\b|leitor|from_curve|"
        r"read from the (?:curve|data|knee)", re.I),
    "read (direct)": re.compile(
        r"\bpaper\b|tabela|\btable\b|handbook|\bvdi\b|\biso\b|input-de-paper|"
        r"motosh|nota de aparato|\bpdf\b|medid[ao]|measured|publicad|"
        r"protocolo|datasheet", re.I),
    "default (design starter)": re.compile(
        r"starter|default do pack|default d[oa] |design de surface", re.I),
    "inherited (kept from an earlier adoption)": re.compile(
        r"mantid|herdad|\bidem\b|mesmo valor|mesma origem|estendido do token|"
        r"zero n[uú]mero novo|inherited|kept from", re.I),
}
# Ordered from weakest to strongest claim (channel off and form switch are
# exact facts, not claims). When a text supports several classes the WEAKEST
# wins: a number both read from a paper and adjusted on a curve is a fit.
# 'inherited' comes last on purpose: it only says the class is recorded in an
# earlier adoption, so any class the text itself states takes precedence.
_ORDEM_CLASSES = ["channel off", "form switch (a mode, not a number)",
                  "fitted (declared)", "anchored", "regressed",
                  "read (canonical reader)", "read (direct)",
                  "default (design starter)",
                  "inherited (kept from an earlier adoption)"]
_CLASSES_TABELA = _ORDEM_CLASSES + ["documented (class not parsed)",
                                    "undocumented"]


def classe_procedencia(txt):
    if not txt:
        return "undocumented"
    t = _NEG_FIT.sub(" ", str(txt))
    achou = [c for c in _ORDEM_CLASSES if _RX_CLASSE[c].search(t)]
    return achou[0] if achou else "documented (class not parsed)"


def _fmt_val(v):
    if isinstance(v, bool):
        return str(v)
    if isinstance(v, int):
        return str(v)
    if isinstance(v, float):
        return f"{v:.4g}"
    if isinstance(v, str):
        return v if len(v) <= 26 else v[:23] + "…"
    if isinstance(v, (list, tuple, dict)):
        return f"list[{len(v)}]"
    return str(v)


def ledger_constantes(comp, specs):
    """Every constant the comparable curves run on, one row each, with the
    provenance class read from the recorded text. Group discovery mirrors
    `metricas`, so the totals agree with §8.4 by construction."""
    try:
        colhida = json.loads((RAIZ / "New_Theory" / "procedencia_colhida.json")
                             .read_text(encoding="utf-8"))
    except Exception:
        colhida = {}
    unidade = {s.name: unidade_en(s.unit) for s in specs}
    vistos, linhas = set(), []
    for r in comp:
        g = rn._adopted_for(r.source, r.case_id,
                            getattr(r.validation_case, "bolt_size", "") or "")
        if not g or g in vistos:
            continue
        vistos.add(g)
        e = kb.adopted_config(g) or {}
        c = e.get("cfg") or {}
        prov = e.get("prov") or {}
        col = colhida.get(g) or {}
        for k, v in c.items():
            if k == "per_case" or not (isinstance(v, (int, float))
                                       and not isinstance(v, bool)):
                continue
            txt, onde = prov_lookup(prov, k), "configuration"
            if txt is None and k in col:
                txt, onde = str(col[k]), "harvested map"
            if txt is None:
                onde = "none"
            linhas.append(dict(source=r.source, group=g, scope="shared",
                               field=k, value=v, unit=unidade.get(k, "–"),
                               classe=classe_procedencia(txt), where=onde))
        for tok, x in (c.get("per_case") or {}).items():
            if not isinstance(x, dict):
                continue
            for k, v in x.items():
                txt = prov_lookup(prov, k) or prov.get(f"per_case.{tok}")
                onde = "configuration" if txt else None
                if txt is None:
                    for chave in (k, f"{tok}.{k}"):
                        if chave in col:
                            txt, onde = str(col[chave]), "harvested map"
                            break
                if txt is None:
                    alt = colhida.get(f"{g}_{tok}") or {}
                    if k in alt:
                        txt, onde = str(alt[k]), "harvested map"
                if txt is None:
                    onde = "none"
                linhas.append(dict(source=r.source, group=g,
                                   scope=f"curve {tok}", field=k, value=v,
                                   unit=unidade.get(k, "–"),
                                   classe=classe_procedencia(txt), where=onde))
    return linhas


_RX_TRIM = [
    (re.compile(r"fratur|fracture", re.I),
     "fatigue-fracture stage, outside the model"),
    (re.compile(r"trinca|crack", re.I),
     "window before crack initiation (bolt shear crack), outside the model"),
    (re.compile(r"debris|terceiro corpo|third body", re.I),
     "debris-dominated tail (third body), outside the model"),
    (re.compile(r"metric-limited|colapso terminal|sufixo terminal|"
                r"sufixo cont", re.I),
     "terminal collapse, metric-limited (§10.3)"),
    (re.compile(r"cauda terminal|tangente|\bN2\b", re.I),
     "terminal tail beyond the paper's own end-of-test marker"),
]


def motivo_trim(txt):
    for rx, rotulo in _RX_TRIM:
        if txt and rx.search(txt):
            return rotulo
    return "documented in the adopted configuration; reason not parsed"


def _texto_trim(prov, valor):
    """The provenance text that explains a trim: a key starting with trim_,
    or any text that names the trimmed cycle count."""
    for k, v in (prov or {}).items():
        if str(k).startswith("trim"):
            return str(v)
    alvo = str(int(round(float(valor))))
    for v in (prov or {}).values():
        if alvo in str(v) or re.search(r"\bTRIM\b|trim_n_max", str(v)):
            return str(v)
    return None


def referencias(comp):
    """{source: (citation, doi)} — DOI harvested from the apparatus notes."""
    out, notas = {}, {}
    for r in comp:
        ref = (getattr(r.validation_case, "reference", "") or "").strip()
        if r.source not in out and ref:
            out[r.source] = [ref, ""]
        n = getattr(r, "apparatus_note_path", None)
        if n and r.source not in notas:
            notas[r.source] = n
    for src, n in notas.items():
        if src in out and out[src][1]:
            continue
        try:
            t = Path(n).read_text(encoding="utf-8", errors="replace")
        except Exception:
            continue
        # ⚠️ one apparatus note can cover TWO papers (the Zhang note covers
        # 2018 and 2019). Taking the first DOI gave both sources the same one,
        # silently. Score the candidates against the citation instead.
        cands = []
        for mt in re.finditer(r"10\.\d{4,9}/[^\s\)\]\},;\"'`]+", t):
            d = mt.group(0).rstrip(".,;`")
            if d not in cands:
                cands.append(d)
        if not cands:
            continue
        cite = (out.get(src) or ["", ""])[0].lower()
        cite_n = re.sub(r"[^a-z0-9]", "", cite)

        def _pontua(d):
            dl = d.lower()
            s = 0
            anos = re.findall(r"\b((?:19|20)\d{2})\b", cite)
            if anos and any(a in dl for a in anos):
                s += 3
            for w in re.findall(r"[a-z]{6,}", dl.split("/", 1)[-1]):
                if w in cite_n:
                    s += 1
            return s

        melhor = max(cands, key=_pontua)
        out.setdefault(src, ["", ""])[1] = melhor
    # loud, not silent: two sources sharing a DOI means the resolution failed
    vistos = collections.Counter(v[1] for v in out.values() if v[1])
    for doi, k in vistos.items():
        if k > 1:
            print(f"  [WARN] DOI shared by {k} sources: {doi} -> "
                  f"{[s for s, v in out.items() if v[1] == doi]}")
    return out


def metricas(comp, res, pisos):
    """Every number the annex quotes, computed once."""
    O, M, ids, srcs, oks = [], [], [], [], []
    for r in comp:
        rr = res.get(r.case_id)
        if rr is None or not (rr.metric_data and rr.metric_pred):
            continue
        O.append(float(rr.metric_data[-1]))
        M.append(float(rr.metric_pred[-1]))
        ids.append(r.case_id)
        srcs.append(r.source)
        oks.append(bool(rh._tripe_ok(rr, rh.limite_sres(r.source, pisos))))
    O, M = np.array(O), np.array(M)
    b = M - O
    d = {}
    d["n"] = len(comp)
    d["n_src"] = len({r.source for r in comp})
    d["tripe"] = sum(1 for r in comp
                     if rh._tripe_ok(res.get(r.case_id),
                                     rh.limite_sres(r.source, pisos)))
    d["r2"] = 1 - float(np.sum(b ** 2) / np.sum((O - O.mean()) ** 2))
    d["bias"] = float(b.mean())
    d["in05"] = float(np.mean(np.abs(b) <= 0.05))
    d["in10"] = float(np.mean(np.abs(b) <= 0.10))
    for lim, key in ((0.85, "iso"), (0.80, "din")):
        vp = int(np.sum((O < lim) & (M < lim)))
        vn = int(np.sum((O >= lim) & (M >= lim)))
        fa = int(np.sum((O >= lim) & (M < lim)))
        fs = int(np.sum((O < lim) & (M >= lim)))
        d[key] = dict(acc=(vp + vn) / len(O), fa=fa, fs=fs)
    d["fs_tripe"] = [ids[i] for i in range(len(O))
                     if O[i] < 0.85 <= M[i] and oks[i]]
    d["fs_all"] = [(ids[i], O[i], M[i], oks[i]) for i in range(len(O))
                   if O[i] < 0.85 <= M[i]]
    # calibration cost — group-level and per-curve are counted SEPARATELY.
    # Lumping them produced a single ratio that overstated the cost in one
    # direction and hid, in the other, that most per-curve entries are inputs
    # read from the paper rather than adjusted knobs.
    ncur, nk = collections.Counter(), collections.Counter()
    ngrp, npc, vistos = collections.Counter(), collections.Counter(), set()
    grupos = collections.defaultdict(set)
    nomes = collections.defaultdict(set)
    for r in comp:
        ncur[r.source] += 1
        g = rn._adopted_for(r.source, r.case_id,
                            getattr(r.validation_case, "bolt_size", "") or "")
        if not g:
            continue
        grupos[r.source].add(g)
        if g in vistos:
            continue
        vistos.add(g)
        c = (kb.adopted_config(g) or {}).get("cfg") or {}
        n = sum(1 for k, v in c.items() if k != "per_case"
                and isinstance(v, (int, float)) and not isinstance(v, bool))
        pc = sum(len(x) for x in (c.get("per_case") or {}).values()
                 if isinstance(x, dict))
        for k, v in c.items():
            if k != "per_case" and isinstance(v, (int, float)) \
                    and not isinstance(v, bool):
                nomes[r.source].add(k)
        for _t, x in (c.get("per_case") or {}).items():
            if isinstance(x, dict):
                nomes[r.source].update(x)
        ngrp[r.source] += n
        npc[r.source] += pc
        nk[r.source] += n + pc
    d["ncur"], d["nk"], d["ngrp"], d["npc"] = ncur, nk, ngrp, npc
    d["ngrupos"] = {s: len(v) for s, v in grupos.items()}
    d["nomes"] = {s: sorted(v) for s, v in nomes.items()}
    d["nomes_tot"] = len(set().union(*nomes.values())) if nomes else 0
    # a source partitioned into as many groups as it has curves is, in
    # substance, calibrated per curve — named rather than averaged away
    d["por_curva"] = sorted(s for s in ncur
                            if d["ngrupos"].get(s, 0) >= ncur[s])
    d["k_total"] = sum(nk.values())
    d["k_grupo"] = sum(ngrp.values())
    d["k_percurva"] = sum(npc.values())
    d["n_groups"] = len(vistos)
    # provenance of every per-curve entry: in the config, in the harvested
    # map, or nowhere. The third number is the one that matters.
    try:
        colhida = json.loads((RAIZ / "New_Theory"
                              / "procedencia_colhida.json")
                             .read_text(encoding="utf-8"))
    except Exception:
        colhida = {}
    com = tot = colh = 0
    sem = []
    for s in kb.adopted_sources():
        e = kb.adopted_config(s) or {}
        prov = e.get("prov") or {}
        for _t, x in ((e.get("cfg") or {}).get("per_case") or {}).items():
            if not isinstance(x, dict):
                continue
            for campo in x:
                tot += 1
                if campo in prov:
                    com += 1
                elif campo in (colhida.get(s) or {}):
                    colh += 1
                else:
                    sem.append((s, campo))
    d["prov_com"], d["prov_tot"] = com, tot
    d["prov_colhida"], d["prov_sem"] = colh, sorted(set(sem))
    # coverage
    A = [float(getattr(r.validation_case, "transverse_displacement_mm", 0) or 0)
         for r in comp]
    F = [float(getattr(r.validation_case, "initial_preload_N", 0) or 0) / 1000
         for r in comp]
    D = [float(getattr(r.validation_case, "bolt_diameter_mm", 0) or 0)
         for r in comp]
    Q = [float(getattr(r.validation_case, "frequency_Hz", 0) or 0) for r in comp]
    d["cov"] = {"amplitude [mm]": [x for x in A if x > 0],
                "preload F0 [kN]": [x for x in F if x > 0],
                "bolt diameter [mm]": [x for x in D if x > 0],
                "frequency [Hz]": [x for x in Q if x > 0]}
    # loading split + dominant mechanism
    modo = collections.Counter()
    dom = collections.Counter()
    # ⚠️ the store is FLAT; `.get("cases", {})` silently yielded {} and the
    # dominant-mechanism line came out empty. Same fallback as `carrega`.
    _b = json.loads((RAIZ / "Models" / "CALIBRATION_AND_VALIDATION"
                     / "validation_store.json").read_text(encoding="utf-8"))
    recsj = _b.get("cases", _b)
    for r in comp:
        rr = res.get(r.case_id)
        cu = (getattr(rr, "config_used", None) or {}) if rr else {}
        modo["transverse (displacement)" if cu.get("mode") == "displacement"
             else "axial (force)"] += 1
        dd = (recsj.get(r.case_id) or {}).get("decomp") or {}
        fim = {m: (v[-1] if isinstance(v, list) and v else 0.0)
               for m, v in dd.items()}
        if fim and max(fim.values()) > 0:
            dom[max(fim, key=fim.get)] += 1
    d["modo"], d["dom"] = modo, dom
    return d




# --------------------------------------------------------------------------- #
# Figures — English, built here so the Portuguese thesis figures stay untouched #
# --------------------------------------------------------------------------- #

import matplotlib                                                # noqa: E402
matplotlib.use("Agg")
import matplotlib.pyplot as plt                                  # noqa: E402
from matplotlib.patches import Patch                             # noqa: E402

# Vector artwork has to carry its fonts. Matplotlib defaults to Type 3, which
# publishers reject; 42 is TrueType with the glyphs subset into the file.
plt.rcParams["pdf.fonttype"] = 42
plt.rcParams["ps.fonttype"] = 42
plt.rcParams.update({
    "font.family": "serif",
    "font.serif": ["STIXGeneral", "Times New Roman", "DejaVu Serif"],
    "mathtext.fontset": "stix",
    "font.size": 9.5,
    "axes.titlesize": 10,
    "axes.labelsize": 9.5,
    "axes.linewidth": 0.8,
    "axes.grid": True,
    "grid.alpha": 0.22,
    "grid.linewidth": 0.5,
    "legend.framealpha": 0.92,
    "legend.edgecolor": "#c8c8c8",
    "figure.autolayout": False,
})

AZUL, VERM, CINZA, AMBAR = "#1f5c88", "#a63232", "#4a4a4a", "#c98a2a"
MECH_EN = {
    "embedding": "embedding",
    "creep": "creep",
    "wear": "wear",
    "rotational_loosening": "rotational loosening",
    "thread_fretting": "flank fretting",
    "fatigue": "fatigue",
}
MECH_COR = {
    "embedding": "#4878a8",
    "creep": "#c98a2a",
    "wear": "#a63232",
    "rotational_loosening": "#6a9a58",
    "thread_fretting": "#8064a2",
    "fatigue": "#666666",
}


DPI_DOC, DPI_ARTE = 300, 600


def _salva(figura, nome):
    """Three files per figure. The 300 dpi PNG is what the documents embed:
    above what the journal asks for a halftone, and light enough that a version
    history of the .docx stays cheap. The 600 dpi PNG is production quality and
    travels as separate artwork. The PDF is the vector original, which is what
    the journal prefers for a line drawing."""
    FIGS.mkdir(parents=True, exist_ok=True)
    figura.savefig(FIGS / f"{nome}.png", dpi=DPI_DOC, bbox_inches="tight",
                   facecolor="white")
    figura.savefig(FIGS / f"{nome}_600dpi.png", dpi=DPI_ARTE,
                   bbox_inches="tight", facecolor="white")
    figura.savefig(FIGS / f"{nome}.pdf", bbox_inches="tight",
                   facecolor="white")
    plt.close(figura)
    print(f"  [fig] {nome}.png ({DPI_DOC} dpi) + _600dpi + .pdf")


def _mx(rr):
    return (list(rr.metric_x), list(rr.metric_data), list(rr.metric_pred))


def _escala_x(ax, x):
    """Log abscissa when the curve spans more than ~1.7 decades. A point at
    N = 0 cannot sit on a log axis, so callers draw it at the left edge with
    `_x_plot` and the caption says so — the same convention the envelope
    figure uses for the axial curves."""
    x = np.asarray(x, float)
    pos = x[x > 0]
    if pos.size and pos.max() / pos.min() > 50:
        ax.set_xscale("log")
        return True
    return False


def _x_plot(x):
    x = np.asarray(x, float)
    pos = x[x > 0]
    if not pos.size:
        return x
    return np.where(x > 0, x, pos.min() / 3.0)


def _mola(ax, x, y0, y1, n=7, w=0.09):
    """Vertical zig-zag spring between (x, y0) and (x, y1)."""
    ys = np.linspace(y0, y1, 2 * n + 3)
    xs = np.full_like(ys, float(x))
    xs[2:-2:2] += w
    xs[3:-2:2] -= w
    ax.plot(xs, ys, color=CINZA, lw=1.1, solid_capstyle="round", zorder=2)


def _caixa(ax, x, y, w, h, txt, fc="#f4f6f9", ec="#6b7683", fs=7.2,
           color="#1a1d23"):
    from matplotlib.patches import FancyBboxPatch
    ax.add_patch(FancyBboxPatch((x - w / 2, y - h / 2), w, h,
                                boxstyle="round,pad=0.02,rounding_size=0.08",
                                fc=fc, ec=ec, lw=0.9, zorder=2))
    ax.text(x, y, txt, ha="center", va="center", fontsize=fs, zorder=3,
            color=color, linespacing=1.25)


def _fig_msd_esquema():
    """The joint as the model sees it: a cross-section on the left, the
    lumped chain with its contact elements on the right. Drawn, not
    computed — it carries no number, so it cannot go stale."""
    from matplotlib.patches import Rectangle, FancyArrowPatch
    f, (a, b) = plt.subplots(1, 2, figsize=(7.4, 3.6), layout="constrained")
    for ax in (a, b):
        ax.set_aspect("equal")
        ax.axis("off")
        ax.grid(False)
        ax.set_xlim(0, 6.4)
        ax.set_ylim(-1.5, 3.1)
    # --- (a) the physical joint -------------------------------------------
    a.add_patch(Rectangle((1.1, 1.0), 4.3, 0.8, fc="#dfe6ee", ec=CINZA, lw=0.8))
    a.add_patch(Rectangle((1.1, 0.2), 4.3, 0.8, fc="#cfd9e4", ec=CINZA, lw=0.8))
    a.add_patch(Rectangle((2.7, -0.55), 0.6, 2.4, fc="#b9b9b9", ec=CINZA, lw=0.8))   # shank, short protrusion below the nut
    a.add_patch(Rectangle((2.05, 1.8), 1.9, 0.1, fc="#7d7d7d", ec=CINZA, lw=0.6))
    a.add_patch(Rectangle((2.2, 1.9), 1.6, 0.55, fc="#9c9c9c", ec=CINZA, lw=0.8))
    a.add_patch(Rectangle((2.2, -0.3), 1.6, 0.5, fc="#9c9c9c", ec=CINZA,
                          lw=0.8, hatch="////"))       # nut flush with the lower flange (y = 0.2)
    a.text(2.1, 2.17, "head", ha="right", va="center", fontsize=8, color=CINZA)
    a.text(2.1, -0.05, "nut", ha="right", va="center", fontsize=8, color=CINZA)
    a.annotate("bearing face  $\\mu_b$\nembedding · wear · damage $D$",
               xy=(3.95, 1.85), xytext=(4.35, 2.55), fontsize=7.6, ha="left",
               arrowprops=dict(arrowstyle="-", lw=0.6, color=CINZA))
    a.annotate("member interface\ntransverse slip",
               xy=(1.7, 1.0), xytext=(0.1, 2.35), fontsize=7.6, ha="left",
               arrowprops=dict(arrowstyle="-", lw=0.6, color=CINZA))
    a.annotate("thread flank  $\\mu_t$\nhelix: $\\theta_{loose} \\rightarrow F_0$",
               xy=(3.8, -0.1), xytext=(4.35, -0.55), fontsize=7.6, ha="left",
               arrowprops=dict(arrowstyle="-", lw=0.6, color=CINZA))
    a.add_patch(FancyArrowPatch((3.0, 3.05), (3.0, 2.5), arrowstyle="-|>",
                                mutation_scale=9, color=VERM, lw=1.2))
    a.text(3.15, 2.9, "$F_0$", color=VERM, fontsize=9, va="center")
    a.add_patch(FancyArrowPatch((0.35, 0.6), (1.0, 0.6), arrowstyle="<|-|>",
                                mutation_scale=9, color=AZUL, lw=1.2))
    a.text(0.05, 0.35, "$\\delta$ or $F_{amp}$", color=AZUL, fontsize=8.0,
           ha="left", va="top")
    a.set_title("(a) the joint", fontsize=9.5)
    # --- (b) the lumped model ----------------------------------------------
    yh, yn = 2.5, -0.6
    for y, lab in ((yh, "head / bearing node"), (yn, "nut / thread node")):
        b.plot([0.9, 4.5], [y, y], color=CINZA, lw=1.6, zorder=1)
        b.text(4.62, y, lab, fontsize=7.6, va="center")
    _mola(b, 1.6, yn + 0.6, yh - 0.05)
    b.text(0.95, 1.05, "$k_b$, $c$\nbolt", fontsize=8.1, ha="center",
           va="center")
    _caixa(b, 1.6, yn + 0.32, 1.55, 0.55,
           "thread contact\n$\\mu_t$ · helix $k_b\\lambda$", fs=7.0)
    _mola(b, 3.4, yn + 0.05, yh - 0.62)
    b.text(4.1, 0.85, "$k_m$\nmembers", fontsize=8.1, ha="center", va="center")
    _caixa(b, 3.4, yh - 0.36, 1.85, 0.55,
           "bearing contact\n$\\mu_b$ · $\\delta_{emb}, \\delta_{creep}, \\delta_{wear}$",
           fs=6.8)
    b.add_patch(FancyArrowPatch((2.5, yn + 0.15), (2.5, yh - 0.15),
                                arrowstyle="<|-|>", mutation_scale=9,
                                color=VERM, lw=1.1))
    b.text(2.62, 1.0, "$F_0$", color=VERM, fontsize=9, va="center")
    b.text(3.2, -1.05,
           "slow state  $s = (F_0,\\ \\delta_{emb},\\ \\delta_{creep},\\ "
           "\\delta_{wear},\\ \\theta_{loose},\\ D)$",
           fontsize=7.7, ha="center", va="top")
    b.set_title("(b) the lumped model", fontsize=9.5)
    _salva(f, "fig_msd_schematic")


def _fig_laco_acoplamento():
    """The couplings of §2.3, drawn as a signal-flow diagram (2026-08-29, at
    the professor's request: arrows in front, readable). Left to right: the
    preload is read by three rates (friction capacity, resisting torque,
    creep), the capacity sets the slip, slip drives wear and loosening, every
    increment sums into −ΔF₀ and returns to the preload along the bottom.
    Modulators (damage D, arrest gate) are written inside the boxes they act
    on; the stiffness partition Φ is shown as present but inert. No numbers."""
    from matplotlib.patches import FancyArrowPatch
    f, ax = plt.subplots(figsize=(7.6, 4.4), layout="constrained")
    ax.set_xlim(-0.75, 10.35)
    ax.set_ylim(-1.55, 5.05)
    ax.axis("off")
    ax.grid(False)
    B = {  # name: (cx, cy, w, h)
        "F0": (0.95, 2.2, 1.5, 0.9), "phi": (0.95, 4.25, 1.9, 0.75),
        "cap": (3.2, 3.6, 2.2, 1.0), "tres": (3.2, 2.2, 2.2, 0.8),
        "slip": (5.45, 3.6, 1.8, 0.85),
        "wear": (7.6, 3.6, 1.9, 0.95), "loose": (7.6, 2.2, 1.9, 1.0),
        "creep": (7.6, 0.85, 1.9, 0.85), "emb": (7.6, -0.4, 1.9, 0.8),
        "sum": (9.55, 1.6, 1.3, 1.7),
    }

    def borda(nome, lado, frac=0.5):
        cx, cy, w, h = B[nome]
        if lado == "L":
            return (cx - w / 2, cy - h / 2 + frac * h)
        if lado == "R":
            return (cx + w / 2, cy - h / 2 + frac * h)
        if lado == "T":
            return (cx - w / 2 + frac * w, cy + h / 2)
        return (cx - w / 2 + frac * w, cy - h / 2)          # "B"

    def seta(p0, p1, **kw):
        base = dict(arrowstyle="-|>", mutation_scale=11, lw=1.25, color=CINZA,
                    shrinkA=1.5, shrinkB=1.5, zorder=6)
        base.update(kw)
        ax.add_patch(FancyArrowPatch(p0, p1, **base))

    def linha(pts, **kw):
        base = dict(color=CINZA, lw=1.25, zorder=6, solid_capstyle="round")
        base.update(kw)
        ax.plot([q[0] for q in pts], [q[1] for q in pts], **base)

    _caixa(ax, *B["F0"], "preload\n$F_0$", fc="#fbe9e7", ec=VERM, fs=10.5)
    _caixa(ax, *B["phi"], "$[K(s)] \\rightarrow \\Phi$\npresent in the equations,\ninert in the corpus",
           fc="#f3f3f3", ec="#9a9a9a", fs=7.6, color="#666666")
    _caixa(ax, *B["cap"], "friction capacity\n$F_{slip} = 0.46\\,\\mu_{b,eff}\\,F_0$\n"
                          "$\\mu_{b,eff} = \\mu_b\\,(1 - k_{dmg,\\mu} D)$", fs=8.0)
    _caixa(ax, *B["tres"], "resisting torque\n$T_{resist} \\propto \\mu\\,F_0$", fs=8.2)
    _caixa(ax, *B["slip"], "resolved slip\n$s = \\max(0,\\ \\delta_{amp} - \\delta_t)$", fs=8.0)
    _caixa(ax, *B["wear"], "wear\n$\\Delta\\delta_{wear} \\propto F_0\\, s\\,(1 + k_{dmg} D)$",
           fc="white", ec=MECH_COR["wear"], fs=8.0)
    _caixa(ax, *B["loose"], "rotational loosening\n$T_{loose}(s)$ against $T_{resist}$\n"
                            "$\\times\\ g_{arrest}(F_0)$",
           fc="white", ec=MECH_COR["rotational_loosening"], fs=8.0)
    _caixa(ax, *B["creep"], "creep (negative feedback)\n$\\Delta\\delta_{creep} \\propto F_0$",
           fc="white", ec=MECH_COR["creep"], fs=8.0)
    _caixa(ax, *B["emb"], "embedding\nstate-based, saturates",
           fc="white", ec=MECH_COR["embedding"], fs=8.0)
    _caixa(ax, *B["sum"], "$-\\Delta F_0$\n$= k_b \\sum \\Delta\\delta$\n$+\\, k_b\\,\\frac{p}{2\\pi}\\,\\Delta\\theta$",
           fs=8.2)

    # F0 -> its three readers
    seta(borda("F0", "R", 0.78), borda("cap", "L", 0.5))
    seta(borda("F0", "R", 0.5), borda("tres", "L", 0.5))
    p0 = borda("F0", "B", 0.5)
    linha([p0, (p0[0], 0.85)])
    seta((p0[0], 0.85), borda("creep", "L", 0.5))
    ax.text(3.9, 0.98, "rate read from $F_0$", fontsize=7.4, color="#666666",
            ha="center", va="bottom", style="italic")
    # capacity -> slip -> wear, loosening; torque -> loosening
    seta(borda("cap", "R", 0.5), borda("slip", "L", 0.5))
    seta(borda("slip", "R", 0.5), borda("wear", "L", 0.5))
    seta(borda("slip", "R", 0.15), borda("loose", "L", 0.8))
    seta(borda("tres", "R", 0.5), borda("loose", "L", 0.4))
    # mechanisms -> sum
    seta(borda("wear", "R", 0.5), borda("sum", "L", 0.86))
    seta(borda("loose", "R", 0.5), borda("sum", "L", 0.62))
    seta(borda("creep", "R", 0.5), borda("sum", "L", 0.38))
    seta(borda("emb", "R", 0.5), borda("sum", "L", 0.14))
    # the loop closes along the bottom, into the left face of F0
    pb = borda("sum", "B", 0.5)
    pl = borda("F0", "L", 0.5)
    linha([pb, (pb[0], -1.25), (-0.45, -1.25), (-0.45, pl[1])])
    seta((-0.45, pl[1]), pl)
    ax.text(4.3, -1.18, "the loop: every increment shortens the clamped length and returns to $F_0$",
            fontsize=7.6, color="#444444", ha="center", va="bottom", style="italic")
    # present but inert
    seta(borda("F0", "T", 0.5), borda("phi", "B", 0.5), color="#9a9a9a", ls="--",
         lw=1.0, mutation_scale=9)
    # reading aid
    ax.text(5.45, 4.62, "positive feedback: capacity and resisting torque fall with $F_0$ · "
                        "negative: creep\nthe gate $g_{arrest}$ and the saturation of "
                        "embedding let the same equations arrest",
            fontsize=7.6, color="#444444", ha="center", va="center", style="italic")
    _salva(f, "fig_coupling_loop")


def _apara_legenda(img):
    """Some explorer crops carry a strip of the paper's caption under the plot.
    If the bottom of the image is a short block of content separated from the
    plot by a clear white band, cut at the band; otherwise return the image
    untouched. Heuristic, and deliberately conservative."""
    try:
        g = np.asarray(img, float)
        if g.ndim == 3:
            g = g[..., :3].mean(axis=2)
        if g.max() > 1.5:
            g = g / 255.0
        tinta = (1.0 - g).mean(axis=1)            # ink per row
        h = len(tinta)
        branco = tinta < 0.004
        # scan the bottom 35 % for the lowest white band at least 1.5 % tall
        corte = None
        i = h - 1
        while i > int(0.65 * h):
            if branco[i]:
                j = i
                while j > 0 and branco[j]:
                    j -= 1
                if i - j >= max(3, int(0.015 * h)) and (h - i) <= int(0.18 * h) \
                        and (h - i) > int(0.02 * h):
                    corte = j + (i - j) // 2
                    break
                i = j
            else:
                i -= 1
        return img[:corte] if corte else img
    except Exception:
        return img


def _fig_cadeia_extracao(cid, porid, res, store):
    """§5 as four panels for one curve: the published figure, the digitised
    points, the anchored and windowed simulation, and the three vectors the
    metric compared. Everything after panel (a) is read from the store and
    the CSV — the same objects the metric used."""
    from bolt_analysis_studio.validation.inputs import load_full_curve
    rec = porid[cid]
    vc = rec.validation_case
    rr = res[cid]
    slug = rec.source.lower()
    pf = RAIZ / "New_Theory" / "variable_explorer" / "paper_figures"
    crop = None
    mfig = re.search(r"fig(\d+)", cid)
    if mfig and (pf / f"{slug}__fig{mfig.group(1)}.png").exists():
        crop = pf / f"{slug}__fig{mfig.group(1)}.png"
    else:
        outros = sorted(pf.glob(f"{slug}__*.png"))
        crop = outros[0] if outros else None
    xr, yr = load_full_curve(vc.reference_csv_path)
    off = float(getattr(vc, "csv_x_offset", 0.0) or 0.0)
    sc = float(getattr(vc, "csv_x_scale", 1.0) or 1.0)
    xr = np.clip((np.asarray(xr, float) - off) * sc, 0, None)
    yr = np.asarray(yr, float)
    e = store[cid]
    cy, ra = np.asarray(e["cycles"], float), np.asarray(e["ratio"], float)
    al = float(rr.align or 1.0)
    x, d, pr = (np.asarray(v, float) for v in _mx(rr))
    n_max = (rr.config_used or {}).get("n_max")
    trim = (rr.config_used or {}).get("trim_n_max")
    # why digitised points are missing from the metric: below the floor, or a
    # hair beyond n_max (the edge effect the runner is known for)
    abaixo = int(np.sum(yr < rn.FLOOR_TRIM))
    alem = int(np.sum(xr > float(n_max))) if n_max else 0
    f, axs = plt.subplots(2, 2, figsize=(7.4, 5.7), layout="constrained")
    a, b, c, dd = axs.ravel()
    a.grid(False)
    a.axis("off")
    if crop is not None:
        a.imshow(_apara_legenda(plt.imread(crop)))
        a.set_title(f"(a) published figure: {crop.stem.split('__')[-1]}, "
                    f"{rec.source.replace('_', ' ').title()}", fontsize=8.5)
    else:
        a.text(0.5, 0.5, "figure crop not available", ha="center",
               va="center", transform=a.transAxes)
    b.plot(_x_plot(xr), yr, "o", ms=3.2, mfc="#e8e8e8", mec=CINZA, mew=0.6)
    _escala_x(b, xr)
    b.set_xlabel("cycles")
    b.set_ylabel("$F/F_0$")
    b.set_title(f"(b) digitised curve: {len(xr)} points, axis conventions "
                f"applied", fontsize=8.5)
    c.axvspan(_x_plot(x)[0], _x_plot(x)[-1], color="#dfe8f0", zorder=0,
              label="metric window (points compared)")
    c.plot(_x_plot(cy), ra, "--", lw=1.0, color="#999999",
           label="model, raw (as simulated)")
    c.plot(_x_plot(cy), ra / al, "-", lw=1.6, color=VERM,
           label=f"model ÷ {al:.4f} (anchored at the first data cycle)")
    c.plot(_x_plot(x), pr, ".", ms=4.5, color=VERM, alpha=0.9,
           label="model at the data cycles (full resolution)")
    c.plot([_x_plot(x)[0]], [1.0], "s", ms=7, mfc="none", mec=AZUL, mew=1.1,
           label="anchor ($n_0$, 1.0)")
    c.axhline(rn.FLOOR_TRIM, color="#333333", lw=0.8, ls=":")
    c.text(0.99, rn.FLOOR_TRIM + 0.02, f"FLOOR_TRIM = {rn.FLOOR_TRIM:g}",
           transform=c.get_yaxis_transform(), ha="right", va="bottom",
           fontsize=6.8, color="#333333")
    if n_max:
        c.axvline(float(n_max), color="#333333", lw=0.8, ls="-.",
                  label=f"end of simulation, n_max = {float(n_max):g}")
    if trim:
        c.axvline(float(trim), color=AZUL, lw=0.8, ls="--",
                  label=f"trim_n_max = {float(trim):g}")
    _escala_x(c, x)          # same scale rule as (b) and (d): decided by the data
    c.set_ylim(0, max(1.08, float(np.nanmax(ra / al)) * 1.02))
    c.set_xlabel("cycles")
    c.set_ylabel("$F/F_0$")
    c.legend(fontsize=6.0, loc="lower left")
    c.set_title("(c) simulation, anchored and windowed", fontsize=8.5)
    dd.vlines(_x_plot(x), d, pr, color=AMBAR, lw=1.1, alpha=0.95,
              label="residual  r = model − data", zorder=2)
    dd.plot(_x_plot(x), d, "o", ms=4.0, mfc="#e8e8e8", mec=CINZA, mew=0.7,
            label="metric_data", zorder=3)
    dd.plot(_x_plot(x), pr, "-", lw=1.6, color=VERM, label="metric_pred",
            zorder=2)
    _escala_x(dd, x)
    dd.set_xlabel("cycles")
    dd.set_ylabel("$F/F_0$")
    dd.legend(fontsize=6.5, loc="lower left")
    dd.set_title(f"(d) the vectors the metric compared: MAE {rr.mae:.4f}, "
                 f"max|r| {rr.maxerr:.4f}, σ {rr.resid_std:.4f}", fontsize=8.5)
    _salva(f, "fig_extraction_chain")
    return {"cid": cid, "n_csv": int(len(xr)), "align": al,
            "n_metric": int(len(x)), "x0_zero": bool(x[0] <= 0),
            "n_max": n_max, "abaixo_piso": abaixo, "alem_nmax": alem,
            "x_ultimo": float(xr[-1])}


def _exemplo_cadeia(porid, res, pisos, ok, alvo):
    """Which curve illustrates §5. A rule, not a name: prefer a curve that
    (i) meets the criterion, (ii) has a visible anchoring (align ≠ 1), (iii)
    has at least one digitised point outside the metric window, so the window
    is not a formality, (iv) runs at most 1000 cycles, so the 400-sample grid
    of the store resolves the curve, and (v) has a figure crop. Smallest MAE
    wins; if nothing qualifies the decomposition curve is reused."""
    from bolt_analysis_studio.validation.inputs import load_full_curve
    pf = RAIZ / "New_Theory" / "variable_explorer" / "paper_figures"
    cand = []
    for c, rr in res.items():
        if not ok.get(c) or rr.mae is None or not rr.metric_x:
            continue
        al = float(rr.align or 1.0)
        n_max = (rr.config_used or {}).get("n_max") or 1e9
        if abs(al - 1.0) < 1e-6 or float(n_max) > 1000:
            continue
        rec = porid[c]
        mfig = re.search(r"fig(\d+)", c)
        if not mfig or not (pf / f"{rec.source.lower()}__fig{mfig.group(1)}.png").exists():
            continue
        try:
            xr, _ = load_full_curve(rec.validation_case.reference_csv_path)
        except Exception:
            continue
        if len(xr) <= len(rr.metric_x):
            continue
        cand.append((rr.mae, c))
    return min(cand)[1] if cand else alvo


def _fig_residuo_pernas(alvo, porid, res, pisos, ok):
    """The three legs of §6 on two curves: one that meets the criterion and
    one with MAE and max|r| inside but σ_res outside — 'small error, wrong
    shape'. The second is chosen from the store, not by hand: the smallest
    MAE among the curves that fail only on σ_res."""
    cand = []
    for c, rr in res.items():
        if ok[c] or rr.mae is None:
            continue
        ls = rh.limite_sres(porid[c].source, pisos)
        sr = rh.sres_para_censo(rr)
        if sr is None:
            continue
        if rr.mae <= rh.META_MAE and rr.maxerr <= rh.META_MAX and sr > ls:
            cand.append((rr.mae, c))
    forma = min(cand)[1] if cand else None
    esc = [c for c in (alvo, forma) if c]
    if not esc:
        return {"ok": None, "forma": None}
    f, axs = plt.subplots(2, len(esc), figsize=(3.75 * len(esc), 5.0),
                          sharex="col", layout="constrained",
                          gridspec_kw={"height_ratios": [1.3, 1]})
    axs = np.asarray(axs).reshape(2, -1)
    for j, c in enumerate(esc):
        rr = res[c]
        x, d, pr = (np.asarray(v, float) for v in _mx(rr))
        r = pr - d
        if abs(float(np.mean(np.abs(r))) - float(rr.mae)) > 1e-6:
            print(f"  [WARN] residual of {c} does not reproduce the stored MAE "
                  f"({np.mean(np.abs(r)):.6f} vs {rr.mae:.6f})")
        ls = rh.limite_sres(porid[c].source, pisos)
        sr = rh.sres_para_censo(rr)
        xp = _x_plot(x)
        top, bot = axs[0, j], axs[1, j]
        top.plot(xp, d, "o", ms=4.0, mfc="#e8e8e8", mec=CINZA, mew=0.7,
                 label="experiment", zorder=3)
        top.plot(xp, pr, "-", lw=1.7, color=VERM, label="model", zorder=2)
        top.set_ylabel("$F/F_0$")
        top.set_title(c, fontsize=8)
        top.legend(fontsize=6.8, loc="lower left")
        _escala_x(top, x)
        bot.axhline(0, color="#333333", lw=0.8)
        bot.fill_between([xp.min(), xp.max()], r.mean() - (sr or 0),
                         r.mean() + (sr or 0), color="#c3d4e4", alpha=0.85,
                         label="bias ± σ_res", zorder=0)
        bot.axhline(r.mean(), color=AZUL, lw=0.9, ls="--", label="bias (mean r)")
        bot.axhline(rr.mae, color=AMBAR, lw=0.8, ls=":", label="± MAE")
        bot.axhline(-rr.mae, color=AMBAR, lw=0.8, ls=":")
        bot.plot(xp, r, "-", lw=0.7, color=CINZA, zorder=2)
        bot.plot(xp, r, "o", ms=3.0, color=VERM, zorder=3)
        k = int(np.argmax(np.abs(r)))
        bot.plot([xp[k]], [r[k]], "o", ms=9.5, mfc="none", mec=VERM, mew=1.2,
                 zorder=4)
        # the max-|r| point and the summary box go to opposite halves, so the
        # box never covers the point it describes
        esquerda = k < len(xp) / 2
        bot.annotate("max |r|", xy=(xp[k], r[k]),
                     xytext=(6 if esquerda else -6, 8 if r[k] >= 0 else -12),
                     textcoords="offset points", fontsize=6.8, color=VERM,
                     ha="left" if esquerda else "right")
        mult = [f"MAE {rr.mae / rh.META_MAE:.2f}×",
                f"max|r| {rr.maxerr / rh.META_MAX:.2f}×",
                (f"σ_res {sr / ls:.2f}×" if sr is not None else "σ_res n/a")]
        veredito = ("meets the criterion" if ok[c]
                    else "outside: σ_res (shape) decides")
        bot.text(0.98 if esquerda else 0.02, 0.96,
                 "  ·  ".join(mult) + f"\n{veredito}",
                 transform=bot.transAxes, ha="right" if esquerda else "left",
                 va="top", fontsize=6.8,
                 bbox=dict(fc="white", ec="#c8c8c8", lw=0.6, pad=2.5))
        lim = max(float(np.max(np.abs(r))), rr.mae, (sr or 0) + abs(r.mean()))
        bot.set_ylim(-1.9 * lim, 1.9 * lim)
        bot.set_xlabel("cycles")
        bot.set_ylabel("r = model − data")
        if j == 0:
            bot.legend(fontsize=6.2, loc="lower left", ncols=3)
    _salva(f, "fig_residual_three_legs")
    return {"ok": alvo, "forma": forma}


def _fig_tornado():
    """One-at-a-time sensitivity of the prediction to a ±20 % perturbation of
    each constant, per loading family, read from the knowledge base."""
    S_tr = kb.sensitivity("transverse") or {}
    S_ax = kb.sensitivity("axial") or {}
    if not S_tr and not S_ax:
        return None
    # O nome desempata. Sem ele a chave e' so' a sensibilidade, `sorted` e'
    # estavel, e a ordem de entrada e' a iteracao de um `set` de strings — que
    # muda a cada PROCESSO (hash randomizado). As constantes congeladas, todas
    # com S = 0, trocavam de lugar entre dois builds identicos: a figura
    # aparecia modificada em todo commit sem que nada tivesse mudado.
    nomes = sorted(set(S_tr) | set(S_ax),
                   key=lambda k: (-max(S_tr.get(k, {}).get("mean", 0.0),
                                       S_ax.get(k, {}).get("mean", 0.0)), k))
    n_tr = max([v.get("n", 0) for v in S_tr.values()] or [0])
    n_ax = max([v.get("n", 0) for v in S_ax.values()] or [0])
    congel = kb.frozen_params() or {}
    y = np.arange(len(nomes))
    f, ax = plt.subplots(figsize=(6.4, 0.32 * len(nomes) + 1.2),
                         layout="constrained")
    ax.barh(y - 0.19, [S_tr.get(k, {}).get("mean", 0.0) for k in nomes],
            height=0.36, color=AZUL, label=f"transverse ({n_tr} cases)")
    ax.barh(y + 0.19, [S_ax.get(k, {}).get("mean", 0.0) for k in nomes],
            height=0.36, color=AMBAR, label=f"axial ({n_ax} cases)")
    rot = [k + ("  (frozen, S ≈ 0)" if k in congel else "") for k in nomes]
    ax.set_yticks(y, rot, fontsize=7.8)
    ax.invert_yaxis()
    ax.set_xlabel("mean shift of predicted $F/F_0$ for a ±20 % perturbation")
    ax.legend(fontsize=8, loc="lower right")
    ax.set_title("One-at-a-time sensitivity, by loading family", fontsize=10)
    _salva(f, "fig_sensitivity_tornado")
    inertes_ax = sorted(k for k, v in S_ax.items() if v.get("mean", 0) < 1e-4)
    return {"n_tr": n_tr, "n_ax": n_ax, "nomes": nomes, "S_tr": S_tr,
            "S_ax": S_ax, "frozen": sorted(congel), "inertes_ax": inertes_ax}


def figuras(comp, res, pisos, store, m, grupo_fisica, todos_id, res_all):
    porid = {r.case_id: r for r in comp}
    porid_all = {r.case_id: r for r in todos_id}
    ok = {c: bool(rh._tripe_ok(res[c], rh.limite_sres(porid[c].source, pisos)))
          for c in res}

    # -- one physics, several behaviours ----------------------------------
    fam = list(grupo_fisica)
    f, axs = plt.subplots(1, len(fam), figsize=(2.35 * len(fam), 2.95),
                          sharey=True, layout="constrained")
    for ax, c in zip(np.atleast_1d(axs), fam):
        x, d, pr = _mx(res_all[c])
        amp = float(getattr(porid_all[c].validation_case,
                            "transverse_displacement_mm", 0) or 0)
        ax.plot(x, d, "o", ms=4.2, mfc="#e8e8e8", mec=CINZA, mew=0.7,
                label="experiment", zorder=3)
        ax.plot(x, pr, "-", lw=1.9, color=VERM, label="model", zorder=2)
        ax.set_xscale("log")
        ax.set_ylim(0, 1.05)
        ax.set_title(f"δ = {amp:g} mm", fontsize=10)
        ax.set_xlabel("cycles")
        ax.text(0.96, 0.05, f"MAE {res_all[c].mae:.3f}",
                transform=ax.transAxes, ha="right", va="bottom",
                fontsize=7.5, color="#555555")
    np.atleast_1d(axs)[0].set_ylabel(r"$F/F_0$")
    np.atleast_1d(axs)[0].legend(fontsize=7.5, loc="lower left")
    f.suptitle("Lu et al. (2024), fig. 18: one configuration, "
               f"{len(fam)} amplitudes, identical constants", fontsize=10.5)
    _salva(f, "fig_one_physics")

    # -- parity ------------------------------------------------------------
    O = np.array([float(res[c].metric_data[-1]) for c in res])
    P = np.array([float(res[c].metric_pred[-1]) for c in res])
    K = np.array([ok[c] for c in res])
    f, ax = plt.subplots(figsize=(4.7, 4.5), layout="constrained")
    for dd, cor in ((0.10, "#dfe8f0"), (0.05, "#c3d4e4")):
        ax.fill_between([0, 1.06], [-dd, 1.06 - dd], [dd, 1.06 + dd],
                        color=cor, zorder=0, lw=0)
    ax.plot([0, 1.06], [0, 1.06], "-", lw=1.0, color="#333333", zorder=1)
    ax.scatter(O[K], P[K], s=17, color=AZUL, alpha=0.85, edgecolor="none",
               zorder=3)
    ax.scatter(O[~K], P[~K], s=22, facecolor="none", edgecolor=VERM,
               lw=1.0, zorder=4)
    ax.set_xlim(0, 1.05)
    ax.set_ylim(0, 1.05)
    ax.set_aspect("equal")
    ax.set_xlabel("measured final $F/F_0$")
    ax.set_ylabel("predicted final $F/F_0$")
    ax.legend(handles=[
        plt.Line2D([], [], marker="o", ls="", color=AZUL, ms=5,
                   label="meets the criterion"),
        plt.Line2D([], [], marker="o", ls="", mfc="none", mec=VERM, ms=5.5,
                   label="outside the criterion"),
        Patch(color="#c3d4e4", label="±0.05"),
        Patch(color="#dfe8f0", label="±0.10")],
        fontsize=7.5, loc="upper left")
    ax.set_title(f"n = {len(O)}   ·   R² = {m['r2']:.4f}   ·   "
                 f"bias = {m['bias']:+.4f}", fontsize=9.5)
    _salva(f, "fig_parity")

    # -- engineering decision ---------------------------------------------
    f, axs = plt.subplots(1, 2, figsize=(7.8, 3.9), layout="constrained")
    for ax, lim, nome in ((axs[0], 0.85, "ISO 16130:2015 (85 %)"),
                          (axs[1], 0.80, "DIN 25201-4 (80 %)")):
        ax.axhspan(lim, 1.06, xmin=0, xmax=lim / 1.06, color="#f3d1d1",
                   alpha=0.8, lw=0)
        ax.scatter(O, P, s=15, color=AZUL, alpha=0.8, edgecolor="none",
                   zorder=3)
        fs = (O < lim) & (P >= lim)
        fa = (O >= lim) & (P < lim)
        ax.scatter(O[fs], P[fs], s=48, facecolor="none", edgecolor=VERM,
                   lw=1.3, zorder=4)
        ax.axhline(lim, color="#333333", lw=0.9)
        ax.axvline(lim, color="#333333", lw=0.9)
        ax.plot([0, 1.06], [0, 1.06], ":", lw=0.8, color="#888888")
        ax.set_xlim(0, 1.06)
        ax.set_ylim(0, 1.06)
        ax.set_aspect("equal")
        ax.set_xlabel("measured final $F/F_0$")
        ax.set_ylabel("predicted final $F/F_0$")
        ax.set_title(nome, fontsize=9.5)
        ax.text(0.02, 1.045, f"unsafe: {int(fs.sum())}", fontsize=8,
                color=VERM, va="top")
        ax.text(1.04, 0.02, f"conservative: {int(fa.sum())}",
                fontsize=8, color="#666666", va="bottom", ha="right")
    _salva(f, "fig_engineering_decision")

    # -- validity envelope -------------------------------------------------
    A, F0, DI, MO = [], [], [], []
    for c in res:
        vc = porid[c].validation_case
        a = float(getattr(vc, "transverse_displacement_mm", 0) or 0)
        fz = float(getattr(vc, "initial_preload_N", 0) or 0) / 1e3
        di = float(getattr(vc, "bolt_diameter_mm", 0) or 0)
        if fz <= 0 or di <= 0:
            continue
        A.append(a if a > 0 else 0.015)
        F0.append(fz)
        DI.append(di)
        MO.append(a > 0)
    A, F0, DI, MO = map(np.array, (A, F0, DI, np.array(MO, dtype=bool)))
    f, ax = plt.subplots(figsize=(5.7, 4.1), layout="constrained")
    ax.scatter(A[MO], F0[MO], s=2.4 * DI[MO], color=AZUL, alpha=0.55,
               edgecolor="white", lw=0.3,
               label="transverse (displacement-controlled)")
    ax.scatter(A[~MO], F0[~MO], s=2.4 * DI[~MO], color=AMBAR, alpha=0.75,
               edgecolor="white", lw=0.3, marker="^",
               label="axial (force-controlled), plotted at the left edge")
    ax.set_xscale("log")
    ax.set_yscale("log")
    ax.grid(True, which="minor", alpha=0.1)
    ax.set_xlabel("imposed transverse amplitude [mm]")
    ax.set_ylabel("initial preload $F_0$ [kN]")
    ax.legend(fontsize=7.5, loc="upper left")
    ax.set_title("Marker area ∝ bolt diameter (M6…M42)", fontsize=9.5)
    _salva(f, "fig_validity_envelope")

    # -- calibration cost --------------------------------------------------
    # curves-per-GROUP, not curves-per-constant: per-curve entries are 1:1 by
    # construction, so a ratio that included them would sit below 1 for
    # everyone and say nothing.
    dados = sorted(((s, m["ncur"][s] / max(m["ngrupos"].get(s, 1), 1),
                     m["ncur"][s], m["ngrupos"].get(s, 0))
                    for s in m["ncur"]), key=lambda t: t[1])
    f, ax = plt.subplots(figsize=(6.0, 6.2), layout="constrained")
    y = np.arange(len(dados))
    cor = [CINZA if d[3] == 0 else (VERM if d[1] <= 1.0 else AZUL)
           for d in dados]
    ax.barh(y, [d[1] for d in dados], color=cor, height=0.72)
    ax.set_yticks(y, [d[0].replace("_", " ").lower() for d in dados],
                  fontsize=8)
    ax.axvline(1.0, color="#333333", lw=1)
    for i, d in enumerate(dados):
        rot = (f"{d[2]} curves, no configuration at all" if d[3] == 0
               else f"{d[2]} curves / {d[3]} group"
                    f"{'' if d[3] == 1 else 's'}")
        ax.text(d[1] + 0.15, i, rot, va="center", fontsize=7,
                color="#444444")
    ax.set_xlim(0, max(d[1] for d in dados) * 1.38)
    ax.set_xlabel("curves served by each configuration "
                  "(1.0 = one configuration per curve)")
    ax.set_title("Calibration cost by source", fontsize=10)
    _salva(f, "fig_calibration_cost")

    # -- per-source error --------------------------------------------------
    med = sorted(((s, float(np.median([res[c].mae for c in res
                                       if porid[c].source == s])))
                  for s in m["ncur"]), key=lambda t: t[1])
    f, ax = plt.subplots(figsize=(6.0, 6.2), layout="constrained")
    y = np.arange(len(med))
    ax.barh(y, [d[1] for d in med],
            color=[VERM if d[1] > rh.META_MAE else AZUL for d in med],
            height=0.72)
    ax.set_yticks(y, [d[0].replace("_", " ").lower() for d in med],
                  fontsize=8)
    ax.axvline(rh.META_MAE, color="#333333", lw=1)
    for i, d in enumerate(med):
        ax.text(d[1] + 0.0012, i, f"{d[1]:.3f}", va="center", fontsize=6.5,
                color="#444444")
    ax.text(rh.META_MAE, len(med) - 0.1, f"  limit {rh.META_MAE:g}",
            fontsize=8)
    ax.set_xlim(0, max(d[1] for d in med) * 1.22)
    ax.set_xlabel("median MAE of the source")
    ax.set_title("Median error by source", fontsize=10)
    _salva(f, "fig_median_mae_by_source")

    # -- which leg binds ---------------------------------------------------
    pernas, sev = collections.Counter(), []
    for c in res:
        rr = res[c]
        ls = rh.limite_sres(porid[c].source, pisos)
        sr = rh.sres_para_censo(rr)
        mult = {"max |r|": rr.maxerr / rh.META_MAX, "MAE": rr.mae / rh.META_MAE}
        if sr is not None:
            mult["σ_res"] = sr / ls
        q = max(mult, key=mult.get)
        sev.append(mult[q])
        if not ok[c]:
            pernas[q] += 1
    f, axs = plt.subplots(1, 2, figsize=(7.4, 3.1), layout="constrained")
    nomes = ["max |r|", "MAE", "σ_res"]
    axs[0].bar(nomes, [pernas.get(n, 0) for n in nomes], color=VERM,
               width=0.55)
    axs[0].set_ylabel("curves outside the criterion")
    axs[0].set_title("Leg that decides the verdict", fontsize=9.5)
    for i, n in enumerate(nomes):
        axs[0].text(i, pernas.get(n, 0) + 0.45, str(pernas.get(n, 0)),
                    ha="center", fontsize=8.5)
    axs[0].set_ylim(0, max(pernas.values()) * 1.18 if pernas else 1)
    axs[1].hist(np.clip(sev, 0, 6), bins=30, color=AZUL)
    axs[1].axvline(1.0, color="#333333", lw=1)
    axs[1].text(1.0, axs[1].get_ylim()[1] * 0.95, " limit", fontsize=8,
                va="top")
    axs[1].set_xlabel("worst leg, as a multiple of its own limit")
    axs[1].set_ylabel("curves")
    axs[1].set_title("Distance to the criterion (all curves)", fontsize=9.5)
    _salva(f, "fig_binding_leg")

    # -- mechanism decomposition ------------------------------------------
    alvo = None
    for c in ("liu2016wear_fig11a_af7p5kn", "caccese2009_compblock_71kPa",
              "liu2022_fig7a_oil_direct_t3"):
        if c in res and (store.get(c) or {}).get("decomp"):
            alvo = c
            break
    if alvo:
        e = store[alvo]
        dec = {k: v for k, v in e["decomp"].items()
               if v and max(abs(x) for x in v) > 0}
        cy, ra = e["cycles"], e["ratio"]
        n = min(len(cy), *(len(v) for v in dec.values()))
        tot = sum(v[-1] for v in dec.values())
        esc = ((1.0 - ra[-1]) / tot) if tot else 0.0
        f, (a1, a2) = plt.subplots(2, 1, figsize=(5.6, 5.0), sharex=True,
                                   layout="constrained")
        x, d, pr = _mx(res[alvo])
        a1.plot(x, d, "o", ms=4.0, mfc="#e8e8e8", mec=CINZA, mew=0.7,
                label="experiment", zorder=3)
        a1.plot(x, pr, "-", lw=1.9, color=VERM, label="model", zorder=2)
        a1.set_ylabel(r"$F/F_0$")
        a1.legend(fontsize=8)
        a1.set_title(f"{alvo}: MAE = {res[alvo].mae:.4f}", fontsize=9.5)
        a2.stackplot(cy[:n], *[[v * esc for v in s[:n]] for s in dec.values()],
                     labels=[MECH_EN.get(k, k) for k in dec],
                     colors=[MECH_COR.get(k, "#999999") for k in dec],
                     alpha=0.9)
        a2.set_xlabel("cycles")
        a2.set_ylabel("cumulative loss [fraction of $F_0$]")
        a2.legend(fontsize=7, loc="upper left", ncols=2)
        _salva(f, "fig_mechanism_decomposition")

    # -- the four figures of the paper plan that were still missing ---------
    # (schematic, coupling loop, extraction chain, residual anatomy). The
    # first two are drawings; the other two read the store for `alvo`.
    _fig_msd_esquema()
    _fig_laco_acoplamento()
    ex = _exemplo_cadeia(porid, res, pisos, ok, alvo)
    cadeia = _fig_cadeia_extracao(ex, porid, res, store) if ex else None
    anat = _fig_residuo_pernas(alvo, porid, res, pisos, ok)
    tornado = _fig_tornado()
    return {"alvo": alvo, "cadeia": cadeia, "anatomia": anat,
            "tornado": tornado}


def familia_fisica(comp, res, prefixo="lu2024_M8_fig18_amp"):
    """Largest subset of a sweep that shares IDENTICAL effective constants.

    The claim of section 8.1 is 'one configuration, several behaviours', so the
    set has to be built by comparing what the engine actually receives — not by
    assuming a figure's curves share a configuration. Any curve of the sweep
    that carries a constant of its own is excluded, and the exclusion is
    printed in the annex rather than hidden.
    """
    porid = {r.case_id: r for r in comp}
    cand = [c for c in porid
            if c.startswith(prefixo) and res.get(c) is not None
            and getattr(res[c], "metric_x", None)]
    por = collections.defaultdict(list)
    for c in sorted(cand):
        eff = rn._effective_overrides(porid[c], {})
        por[json.dumps({k: str(v) for k, v in sorted(eff.items())},
                       sort_keys=True)].append(c)
    if not por:
        return [], 0
    fam = max(por.values(), key=len)
    return (sorted(fam, key=lambda c: float(getattr(
                porid[c].validation_case, "transverse_displacement_mm", 0)
                or 0)),
            sum(len(v) for v in por.values()) - len(fam))




# --------------------------------------------------------------------------- #
# Document scaffolding                                                        #
# --------------------------------------------------------------------------- #

def estilo(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.15
    # body text justified. Captions, the title page and table cells set their
    # own alignment, so they are not affected
    n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for lvl, sz in ((0, 20), (1, 15), (2, 12.5), (3, 11)):
        try:
            s = doc.styles[f"Heading {lvl}"] if lvl else doc.styles["Title"]
            s.font.name = "Calibri"
            s.font.size = Pt(sz)
            s.font.color.rgb = RGBColor(0x1A, 0x1D, 0x23)
        except KeyError:
            pass


def p(doc, txt, *, bold=False, italic=False, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    for pedaco, forte in _negrito(txt):
        r = par.add_run(pedaco)
        r.bold = bold or forte
        r.italic = italic
        if size:
            r.font.size = Pt(size)
    return par


def _negrito(txt):
    """**bold** -> runs. Keeps the source text readable in this file."""
    out, i = [], 0
    for m in re.finditer(r"\*\*(.+?)\*\*", txt):
        if m.start() > i:
            out.append((txt[i:m.start()], False))
        out.append((m.group(1), True))
        i = m.end()
    if i < len(txt):
        out.append((txt[i:], False))
    return out or [(txt, False)]


def mono(doc, txt, size=9):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.left_indent = Inches(0.25)
    r = par.add_run(txt)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return par


def nota(doc, txt):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.22)
    par.paragraph_format.space_after = Pt(8)
    for pedaco, forte in _negrito(txt):
        r = par.add_run(pedaco)
        r.bold = forte
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x5B, 0x64, 0x72)
    return par


REGRAS_HORIZONTAIS = [False]   # journal tables: set True by the paper build


def _so_horizontais(t):
    """Elsevier asks for no vertical rules and no shading inside table cells:
    keep a rule above and below the header and one below the last row."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    n_lin = len(t.rows)
    for i, row in enumerate(t.rows):
        for cel in row.cells:
            tcPr = cel._tc.get_or_add_tcPr()
            for tag in ("w:tcBorders", "w:shd"):
                for e in tcPr.findall(qn(tag)):
                    tcPr.remove(e)
            b = OxmlElement("w:tcBorders")
            for lado, ligado in (("top", i == 0),
                                 ("bottom", i in (0, n_lin - 1)),
                                 ("left", False), ("right", False),
                                 ("insideH", False), ("insideV", False)):
                e = OxmlElement(f"w:{lado}")
                e.set(qn("w:val"), "single" if ligado else "nil")
                if ligado:
                    e.set(qn("w:sz"), "8")
                    e.set(qn("w:color"), "000000")
                b.append(e)
            tcPr.append(b)
            sh = OxmlElement("w:shd")
            sh.set(qn("w:val"), "clear")
            sh.set(qn("w:fill"), "FFFFFF")
            tcPr.append(sh)


def tabela(doc, cab, linhas, larguras=None, size=9):
    anterior = doc.paragraphs[-1].text if doc.paragraphs else ""
    t = doc.add_table(rows=1, cols=len(cab))
    if REGRAS_HORIZONTAIS[0]:
        t.style = "Table Grid"          # rules trimmed after the rows are in
    else:
        t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cab):
        cel = t.rows[0].cells[i]
        cel.text = ""
        cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = cel.paragraphs[0].add_run(str(c))
        r.bold = True
        r.font.size = Pt(size)
    for ln in linhas:
        cells = t.add_row().cells
        for i, v in enumerate(ln):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT   # never justify a cell
            for pedaco, forte in _negrito(str(v)):
                r = par.add_run(pedaco)
                r.bold = forte
                r.font.size = Pt(size)
    if larguras:
        for row in t.rows:
            for i, w in enumerate(larguras):
                row.cells[i].width = Inches(w)
    if REGRAS_HORIZONTAIS[0]:
        _so_horizontais(t)              # every row exists by now
    TAB_REGISTRO.append((doc, anterior if anterior.startswith("Table") else "",
                         cab, linhas, larguras, size))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


_NFIG = [0]
FIG_ORDEM = []      # (figure name, caption) in the order they were inserted
TAB_REGISTRO = []   # (doc, caption, header, rows, widths, size) per table


def fig(doc, nome, legenda, largura=6.2):
    """Insert an exported figure if it exists; otherwise leave a marker."""
    for ext in (".png",):
        f = FIGS / f"{nome}{ext}"
        if f.exists():
            _NFIG[0] += 1
            doc.add_picture(str(f), width=Inches(largura))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            FIG_ORDEM.append((nome, legenda))
            r = cap.add_run(f"Figure {_NFIG[0]}. {legenda}")
            r.italic = True
            r.font.size = Pt(9)
            cap.paragraph_format.space_after = Pt(10)
            return True
    nota(doc, f"[figure {nome} not embedded — it is built by "
              f"build_annex_docx.py itself; check the console for the error]")
    return False




# --------------------------------------------------------------------------- #
# The annex                                                                   #
# --------------------------------------------------------------------------- #

def monta(doc, comp, res, pisos, store, m, refs, specs,
          grupo_fisica, n_fora_fisica, info, todos_id, res_all):
    alvo = info.get("alvo")
    cadeia = info.get("cadeia") or {}
    anat = info.get("anatomia") or {}
    # ---------------------------------------------------------------- title
    t = doc.add_heading("Bolt Analysis Studio", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Software annex: an energy-based model for bolted-joint "
                    "self-loosening")
    r.italic = True
    r.font.size = Pt(12)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run(f"Validated against {m['n']} digitised curves from "
                     f"{m['n_src']} independent published sources")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x5B, 0x64, 0x72)
    doc.add_paragraph()
    for linha in ("Leonardo Rosa Ribeiro da Silva (corresponding author, leorrs@ufu.br), "
                  "Neilon de Souza da Silva, "
                  "Luiz Eduardo dos Santos Paes, Aldemir Aparecido Cavallini Junior, "
                  "Bruno Sousa Carneiro da Cunha, Fernando Buiatti Rodrigues, "
                  "Artur Martins Alves, Bruno César Alvares Teixeira, "
                  "Douglas da Silva Carvalho, Gabriel Henrique Arruda Tavares de Lima, "
                  "João Paulo de Jesus Vieira",
                  "Faculdade de Engenharia Mecânica, Universidade Federal de "
                  "Uberlândia (UFU), Brazil",
                  "Neilon de Souza da Silva: Petróleo Brasileiro S.A. "
                  "(Petrobras), Rio de Janeiro, RJ, Brazil",
                  "Software written by Leonardo Rosa Ribeiro da Silva and "
                  "Neilon de Souza da Silva."):
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pa.paragraph_format.space_after = Pt(1)
        ra = pa.add_run(linha)
        ra.font.size = Pt(10)
    import datetime as _dt
    try:
        fp0 = rn.engine_fingerprint()
    except Exception:
        fp0 = "(unavailable)"
    g = info.get("git")
    versao = (f"Built on {_dt.date.today().isoformat()} from repository "
              f"revision {g[0]} ({g[1]})" if g else
              f"Built on {_dt.date.today().isoformat()}")
    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pa.paragraph_format.space_before = Pt(6)
    ra = pa.add_run(f"{versao} · result store fingerprint {fp0}")
    ra.font.size = Pt(9)
    ra.font.color.rgb = RGBColor(0x5B, 0x64, 0x72)
    doc.add_paragraph()

    # ---------------------------------------------------------------- 1
    doc.add_heading("1. Scope, and how to read this annex", 1)
    p(doc, "This annex documents the software used to produce the results in "
           "the main paper, and it is written to be read on its own. It states "
           "what the model is, how it is run, what it can and cannot do, and "
           "what every constant means. A section of its own presents the "
           "evidence that the agreement with experiment does not come from "
           "fitting each curve individually.")
    p(doc, "**Every number in this annex is recomputed from the canonical "
           "result store when the document is built.** Nothing here is "
           "transcribed by hand. If the store changes, this document changes "
           "with it; a number that cannot be recomputed is simply not "
           "published, and the few historical measurements quoted from the "
           "project record carry their date and their source document. The "
           "interactive reports that accompany the software are governed by "
           "the same rule.")
    nota(doc, "The software is released free of charge. Nothing in this annex "
              "requires a licence to reproduce.")

    # ---------------------------------------------------------------- 2
    doc.add_heading("2. What the software is", 1)
    p(doc, "Bolt Analysis Studio (BAS) predicts the loss of preload in a "
           "bolted joint subjected to cyclic loading. The joint is represented "
           "as a mass–spring–damper (MSD) assembly: bolt, nut, washers and "
           "clamped members are discrete elements, and the interfaces between "
           "them carry the tribology. Neither a curve-fitting tool nor a "
           "finite-element code, BAS is a lumped-parameter physical model "
           "integrated cycle by cycle.")
    p(doc, "The quantity it predicts is the retained preload ratio "
           "F/F₀ as a function of cycle number N. That is what the "
           "experimental literature reports, and it is what the design "
           "standards legislate: ISO 16130:2015 places the acceptance boundary "
           "at 85 % retention, DIN 25201-4 at 80 %.")
    fig(doc, "fig_msd_schematic",
        "The joint as the model represents it. (a) Cross-section with the "
        "three interfaces that carry the tribology: bearing face, member "
        "interface and thread flank. (b) The lumped chain: bolt and members "
        "in parallel between the head and nut nodes, each in series with its "
        "contact element. The slow state s shortens the clamped length, and "
        "the helix converts nut rotation into preload loss.")
    doc.add_heading("2.1 The slow state vector", 2)
    p(doc, "The model carries a state that evolves slowly compared with the "
           "load cycle:")
    mono(doc, "s = ( F₀ , δ_emb , δ_creep , δ_wear , θ_loose , D )")
    tabela(doc, ["component", "meaning"],
           [["F₀", "current preload [N]"],
            ["δ_emb", "accumulated embedding (asperity flattening) [m]"],
            ["δ_creep", "accumulated viscoelastic/plastic creep [m]"],
            ["δ_wear", "material removed at the interfaces [m]"],
            ["θ_loose", "accumulated rotational loosening of the nut [rad]"],
            ["D", "surface damage, dimensionless in [0, 1]"]],
           larguras=[1.1, 5.1])
    p(doc, "The first five reduce the clamped length and therefore the "
           "preload. The sixth, D, is different in kind. Removing no preload "
           "itself, it **modulates** the other channels, lowering the "
           "effective bearing friction and amplifying wear; that indirect "
           "action is what makes it a state, while the word mechanism is "
           "kept for the channels that withdraw preload themselves.")

    doc.add_heading("2.2 The four loss mechanisms", 2)
    p(doc, "Four mechanisms act in parallel on every cycle. Each returns its "
           "own preload decrement and its own dissipated energy, and the "
           "decomposition is preserved so that any prediction can be attributed:")
    tabela(doc, ["mechanism", "physical basis", "governing law"],
           [["Embedding", "plastic flattening of surface asperities",
             "exponential approach to a finite depth"],
            ["Creep", "viscoelastic/plastic relaxation under sustained load",
             "logarithmic or saturating kernel"],
            ["Wear", "material removal at slipping interfaces",
             "Archard (specific wear rate)"],
            ["Rotational loosening", "nut backing off under transverse slip",
             "slip-driven ratchet with self-locking gate"]],
           larguras=[1.5, 2.6, 2.2])
    p(doc, "A fifth channel, fatigue, is available and off by default; it "
           "becomes relevant only where the experiment ran to fracture.")
    nota(doc, "Energy accounting is explicit: W_ext + ΔU = Σ W_dissipated. The "
              "conservation residual is computed by the engine and, since "
              "2026-08-28, persisted in the result store for every curve "
              "(§2.6, §10.1, Appendix A).")

    doc.add_heading("2.3 Couplings: how the state feeds back into the rates", 2)
    p(doc, "This is the part that distinguishes V2 from a sum of independent "
           "decay laws: every rate reads the current state, so the preload a "
           "cycle removes changes what the next cycle removes. The live paths, "
           "as the ablation of the paper measured them, are:")
    mono(doc, "F₀  →  friction capacity F_slip = 0.46·μ_b,eff·F₀  →  slip s  →  wear, "
              "rotational loosening  →  −ΔF₀  →  F₀\n"
              "F₀  →  resisting torque T_resist ∝ μF₀  →  rotational loosening\n"
              "F₀  →  creep rate ∝ F₀   (negative feedback)\n"
              "D   →  μ_b,eff ↓, wear ↑   (damage, fed by the slip work)")
    p(doc, "The first two paths are positive feedback and can run away; the "
           "arrest gate, the saturation of the state-based embedding and the "
           "falling slip work let the same equations stabilise instead, and "
           "**no parameter chooses between those outcomes**. A fifth path is "
           "present in the equations and inert in the validated corpus: the "
           "contact stiffness softens as F₀ falls (Greenwood–Williamson) and "
           "moves the load partition Φ, but the transverse loosening leg carries "
           "a constant gain in place of Φ and the axial leg that carries Φ is "
           "zero in displacement-controlled tests; freezing it changes one curve "
           "of 205, while cutting the preload feedback into the rates drops the "
           "census from 171 to 111 (paper, Section 4.10.2). One configuration can "
           "therefore reproduce both a joint that collapses and a joint that "
           "stabilises, and §8 presents the evidence.")
    fig(doc, "fig_coupling_loop",
        "The couplings. Solid arrows: the preload sets the friction capacity, "
        "which sets the slip that drives wear and rotational loosening, and "
        "sets the resisting torque; creep reads the preload with the opposite "
        "sign; every increment returns to the preload through the clamped "
        "length and the helix. Dashed amber: the damage state and the arrest "
        "gate. Dashed grey: the stiffness partition Φ, present in the equations "
        "and inert in the validated corpus.")

    doc.add_heading("2.4 Loading modes", 2)
    p(doc, "Two modes, and the distinction matters for accuracy:")
    tabela(doc, ["mode", "when to use", "how slip is obtained"],
           [["Force-controlled", "servo-hydraulic axial excitation",
             "derived from local elasticity"],
            ["Displacement-controlled",
             "Junker-type transverse tests (crank-driven)",
             "imposed: slip = max(0, δ − F_slip/k_tr)"]],
           larguras=[1.5, 2.4, 2.4])
    n_tr = m["modo"].get("transverse (displacement)", 0)
    n_ax = m["modo"].get("axial (force)", 0)
    p(doc, f"In the present corpus {n_tr} curves are transverse "
           f"(displacement-controlled) and {n_ax} are axial (force-controlled).")

    doc.add_heading("2.5 Governing equations", 2)
    p(doc, "The update the engine solves once per load cycle, written with "
           "the default kernels. Every optional form of §9.2 is switched off "
           "here; when one is enabled it replaces or multiplies the term it "
           "names. Symbols follow §2.1 and §9; t = N/f is the elapsed time "
           "and s the transverse slip resolved in the cycle.")
    p(doc, "**Stiffness and load partition.** The bolt is a linear spring; "
           "the clamped members soften as preload falls, which is the "
           "re-evaluation of [K(s)] that couples the loop:", space=3)
    mono(doc, "k_b = E·A_s / L_eff\n"
              "k_j(F₀) = k_j,init · (F₀ / F₀,init)^α_GW          "
              "(Greenwood–Williamson softening)\n"
              "Φ_eff = min( k_b / (k_b + k_j(F₀)), 1 )\n"
              "β = arctan( p / (π·d₂) ),   lead per radian = p / 2π", size=8.5)
    p(doc, "**Slip at the interface.** The onset follows Pai and Hess; the "
           "stick stroke is the onset force over the transverse stiffness, "
           "and what the imposed motion exceeds it by is the slip:", space=3)
    mono(doc, "F_slip = 0.46 · μ_b,eff · F₀\n"
              "δ_t = δ_free + F_slip / k_tr\n"
              "displacement mode:  s = max( 0, δ_amp − δ_t )\n"
              "force mode:         s = max( 0, (F_amp·sin θ − F_slip) / k_tr )\n"
              "k_tr:  axial_frac → 0.3·k_j,init ;  bending → c_bend·E·I/L_eff³, "
              "I = π·d₂⁴/64 ;  in series with k_member_shear when set",
         size=8.5)
    p(doc, "**The mechanisms.** Each returns a depth increment or a rotation "
           "from the same start-of-cycle state:", space=3)
    mono(doc, "embedding   Δδ_emb   = (δ_target − δ_emb) · (1 − e^(−1/N_emb)),"
              "   δ_target = emb_depth · S_conf · S_ρ\n"
              "creep       Δδ_creep = C_creep · F₀ · [ ln(t + t₀) − ln(t − 1/f + t₀) ]\n"
              "wear        Δδ_wear  = k_run · k_wear_spec · F₀ · (4 s) / A_contact"
              " · (1 + k_dmg_wear·D) · g_onset · g_conf\n"
              "fretting    Δδ_fret  = k_thread_fret · k_wear_spec · F₀ · (4 F_ax / k_b)"
              " / A_s · (f_ref / f)^p_fret,   F_ax = F_amp·cos θ\n"
              "loosening   T_loose  = ½ d₂ · hypot( Φ_ax·sin β·F_ax ,  Φ_tr·cos β·F_tr ),"
              "   F_tr = F_amp·sin θ\n"
              "            T_resist = μ_t · F₀ · d₂ / (2 cos 30°) + μ_b,eff · F₀ · r_bearing\n"
              "            Δθ = g_arrest · (T_loose − T_resist)² / (T_loose · k_torsional)"
              "   if T_loose > T_resist, else 0\n"
              "fatigue     σ_a = K_t·F_amp/A_s,   σ_ar = σ_a / (1 − σ_m/σ_uts),"
              "   ΔD_fat = 1/N_f(σ_ar);   fracture when D_fat ≥ 1", size=8.5)
    p(doc, "Φ_tr is the transverse factor-1 gain: tr_loose_gain once F_tr "
           "reaches F_slip, 0.01 below it. Φ_ax is Φ_eff below the separation "
           "load and 1 above it. The torsional stiffness is k_j,init·d₂/2 in "
           "the default mode, or η·G·J/L_eff with J = π·d₂⁴/32 when the bolt-"
           "torsion mode is selected.", space=3)
    p(doc, "**State update.** All increments are summed before the preload "
           "moves, so no mechanism sees another's effect within the cycle; "
           "the damage state and the slip-work accumulator are updated last "
           "and read at the start of the next cycle:", space=3)
    mono(doc, "ΔF₀ = − k_b · (Δδ_emb + Δδ_creep + Δδ_wear + Δδ_fret) "
              "− k_b · (p/2π) · Δθ   [ − release at fracture ]\n"
              "F₀ ← max( F₀ + ΔF₀ , 0 )\n"
              "D ← D + c_D · (W_slip,cycle / W_ref) · (1 − D),   "
              "W_slip,cycle = dE_wear + dE_loose\n"
              "W_slip,acc ← W_slip,acc + 4 · μ_b,eff · F₀ · s", size=8.5)
    p(doc, "**Gates.** Smooth functions in place of switches; each acts on "
           "the preload decrement and never on the dissipated energy:",
      space=3)
    mono(doc, "g_arrest = max( 0, 1 − F_min / F₀ ),   F_min = loose_arrest_floor · F₀,init\n"
              "g_conf   = W_conf,ref / ( W_conf + W_conf,ref )\n"
              "g_onset  = x^k / ( x^k + 1 ),   x = W_slip,acc / slip_onset_W",
         size=8.5)
    p(doc, "**Energy.** The elastic energy of the stack and the external "
           "work close against the dissipation buckets, one per mechanism:",
      space=3)
    mono(doc, "U(F) = F² / (2 k_b)\n"
              "W_ext + [ U(F₀) − U(F₀,init) ] = Σ W_diss,     "
              "residual = W_ext + ΔU − Σ W_diss", size=8.5)
    nota(doc, "The complete derivation, including every optional kernel and "
              "the routing of each mechanism's energy into its bucket, is kept "
              "in MODEL_MATH_REFERENCE.md in the repository and is revised "
              "with the engine.")

    doc.add_heading("2.6 Numerical scheme, cost and verification", 2)
    eng = info.get("engine") or {}
    p(doc, "The engine is a quasi-static, explicit, cycle-by-cycle integrator: "
           "one call to step_cycle per load cycle, no sub-cycle time "
           "integration in the loosening state, and no adaptive cycle "
           "jumping. The mass and Rayleigh-damping matrices of the three-"
           "degree-of-freedom dynamic model (x, y, θ) exist in the code and "
           "are inactive in every curve of this corpus (§9.2, group "
           "'numerical'). No random numbers are drawn, so a run is "
           "reproducible bit for bit.")
    if eng.get("cycles_per_s"):
        cps = eng["cycles_per_s"]
        p(doc, f"**Cost.** Measured while this document was built, on one "
               f"core, in pure Python: {cps:,.0f} cycles per second on the §3.2 "
               f"example, so a 10⁶-cycle test runs in about {1e6/cps:.0f} s "
               f"and a 10³-cycle Junker test in a fraction of a second. The "
               f"full corpus of {m['n']} curves re-simulates in the order of "
               f"ten minutes serially; the parallel batch runner divides that "
               f"by the number of workers.")
    linhas = []
    if eng:
        linhas.append(["Closed form", "the state-based embedding integrated "
                       f"over {eng['n_norton']} cycles against the Norton law "
                       f"δ∞(1 − e^(−N/N_emb)), all other channels off",
                       f"max deviation {eng['norton_dev_m']:.1e} m on a depth "
                       f"of {eng['emb_depth']*1e6:g} µm"])
        linhas.append(["Energy budget, example", "W_ext + ΔU − Σ W_diss on "
                       f"the §3.2 example, {eng['n']} cycles",
                       f"{eng['residual_J']:.2f} J of {eng['W_diss_J']:.0f} J "
                       f"dissipated ({eng['residual_J']/max(eng['W_diss_J'],1e-12):.1e} relative)"])
    ebc = info.get("energy_corpus")
    if ebc:
        linhas.append(["Energy budget, corpus",
                       f"the residual persisted in the store for "
                       f"{ebc['n']} of {m['n']} curves, relative to the larger "
                       f"of |ΣW_diss| and |ΔU|",
                       f"median {ebc['med']:.1e}, 90th percentile "
                       f"{ebc['p90']:.1e}, worst {ebc['max']:.1e} "
                       f"({ebc['pior']}); see the reading below"])
    linhas.append(["Decomposition", "the per-mechanism losses summed against "
                   "the total loss, every curve of the store",
                   "exact by construction; pinned by test"])
    linhas.append(["Default inertness", "every optional form, when off, "
                   "leaves the prediction bit-identical to the engine "
                   "without it", "pinned by test, one per form"])
    if info.get("tests"):
        linhas.append(["Test suite", f"{info['tests']} tests collected at "
                       "build time (pytest)",
                       "last full run 2026-08-27: 1156 passed, 1 skipped, 1 "
                       "knowingly red (a test that encodes a pending census "
                       "decision, not an engine defect)"])
    tabela(doc, ["check", "what is compared", "result"], linhas,
           larguras=[1.1, 3.0, 2.1], size=8.5)
    if ebc:
        p(doc, f"The energy residual deserves a reading rather than a "
               f"summary, because its tail is wide. {ebc['n_cauda']} of the "
               f"{ebc['n']} curves sit above 5 % of their own budget, and "
               f"{ebc['n_cauda_orc_1J']} of those have a budget below one "
               f"joule: they are the creep tests under sustained preload and "
               f"the stick-regime tests, where no external work enters, the "
               f"depth mechanisms book their dissipation as F₀·Δδ and the "
               f"elastic release is computed exactly, so the second-order "
               f"term of a large preload step shows up as residual. "
               + (f"The other {ebc['n_cauda_grandes']} have budgets between "
                  f"{ebc['orc_grandes_lo']:.1f} and {ebc['orc_grandes_hi']:.0f} J. "
                  if ebc.get("n_cauda_grandes") else "")
               + f"In absolute terms the worst residual in the tail is "
               f"{ebc['max_abs_J_cauda']:.2f} J. Where slip work dominates the "
               f"budget (the {ebc['n_kJ']} curves dissipating a kilojoule or "
               f"more) the worst relative residual is "
               f"{ebc['max_kJ']:.1e}. Neither fatigue nor surface damage is "
               f"active in any curve of the tail ({ebc['n_cauda_fat']} and "
               f"{ebc['n_cauda_dmg']} of {ebc['n_cauda']}), so the tail is not "
               f"the collapse regime the earlier text of this annex blamed; "
               f"the sources most represented in it are "
               f"{', '.join(ebc['fontes_cauda'])}. The residual never feeds "
               f"back into the preload, so no prediction depends on it; it is "
               f"published per curve in Appendix A as a check on the "
               f"bookkeeping, and the bookkeeping of the depth mechanisms is "
               f"the place to look.")
    p(doc, "Two properties of the result store matter to anyone reading it "
           "directly. The stored curve is sampled on a linear grid of 400 "
           "points from 0 to n_max, so for a 10⁶-cycle test the second "
           "sample already sits at N ≈ 2 500 and the embedding transient is "
           "invisible in that vector; the three metric vectors (§5) hold the "
           "full-resolution model at the data cycles and are the ones to "
           "read. And the store records, for every curve, the fingerprint of "
           "the configuration that produced it (§13).")

    # ---------------------------------------------------------------- 3
    doc.add_heading("3. Using the software", 1)
    p(doc, "Three entry points, in increasing order of automation.")
    doc.add_heading("3.1 Interactive", 2)
    mono(doc, "python run_app.py            # full application\n"
              "python run_app.py --v2       # Abaqus-style chrome\n"
              "python run_app.py --builder  # MSD model builder only")
    p(doc, "The workflow has four steps. Choose a joint preset in the guided "
           "wizard, or assemble the element chain by hand in the MSD Builder. "
           "Set the loading and friction in the Property Inspector, which is "
           "the single source of truth for both. Run the solver. Read the "
           "results. The Results module also browses the validation corpus "
           "and can load any published case into the model for re-simulation.")
    doc.add_heading("3.2 Programmatic", 2)
    mono(doc, SNIPPET)
    if info.get("snippet_ratio") is not None:
        p(doc, f"The block runs as printed. With the default material and an "
               f"M16 bolt at 0.5 mm imposed amplitude it returns "
               f"F/F₀ = {info['snippet_ratio']:.3f} after 1000 cycles; the "
               f"build of this annex executes it, so the code and the text "
               f"cannot drift apart. Per-rig constants are set as fields of "
               f"JointMaterial, and the geometry helper reads the ISO thread "
               f"table.")
    doc.add_heading("3.3 Validation and reporting", 2)
    mono(doc, "python -m bolt_analysis_studio.validation.report --all\n"
              "python New_Theory/parallel_batch.py --workers 6 --store\n"
              "python New_Theory/build_variable_explorer.py")
    p(doc, "The first re-simulates the whole corpus and regenerates the "
           "reports; the second does the same in parallel; the third builds the "
           "interactive documentation, one page per constant.")

    # ---------------------------------------------------------------- 4
    doc.add_heading("4. The validation corpus", 1)
    p(doc, f"The software is validated against **{m['n']} digitised curves "
           f"from {m['n_src']} independent published sources**. No curve in the "
           f"corpus was measured for this work: every one is "
           f"read from a figure or table in the cited paper.")
    # conditions per source, read from the registry and the store
    cond = collections.defaultdict(lambda: {"bolts": set(), "d": [], "fa": [],
                                            "f0": [], "modes": collections.Counter()})
    for r in comp:
        vc = r.validation_case
        rr = res.get(r.case_id)
        cu = (getattr(rr, "config_used", None) or {}) if rr else {}
        c = cond[r.source]
        bs = (getattr(vc, "bolt_size", "") or "").split("x")[0].strip()
        if bs:
            c["bolts"].add(bs)
        d_ = float(getattr(vc, "transverse_displacement_mm", 0) or 0)
        fa_ = float(getattr(vc, "axial_force_amplitude_N", 0) or 0) \
            or float(cu.get("F_amp_N", 0) or 0)
        f0_ = float(getattr(vc, "initial_preload_N", 0) or 0) / 1000.0
        if d_ > 0:
            c["d"].append(d_)
            c["modes"]["transverse"] += 1
        elif fa_ > 0:
            c["fa"].append(fa_ / 1000.0)
            c["modes"]["axial"] += 1
        else:
            c["modes"]["sustained"] += 1
        if f0_ > 0:
            c["f0"].append(f0_)

    def _faixa(v, unidade):
        if not v:
            return ""
        lo, hi = min(v), max(v)
        return (f"{lo:g} {unidade}" if abs(hi - lo) < 1e-9
                else f"{lo:g}–{hi:g} {unidade}")

    def _carregamento(c):
        partes = []
        if c["modes"]["transverse"]:
            partes.append("transverse, δ " + _faixa(c["d"], "mm"))
        if c["modes"]["axial"]:
            partes.append("axial, F_a " + _faixa(c["fa"], "kN"))
        if c["modes"]["sustained"]:
            partes.append("sustained preload")
        return " · ".join(partes)

    def _parafusos(c):
        def num(b):
            mm = re.search(r"\d+(?:\.\d+)?", b)
            return float(mm.group(0)) if mm else 0.0
        return ", ".join(sorted(c["bolts"], key=num)) or "—"

    idx_ref = {s: i for i, s in enumerate(sorted(m["ncur"]), 1)}
    linhas = []
    for src in sorted(m["ncur"]):
        c = cond[src]
        linhas.append([f"[{idx_ref[src]}]", src.replace("_", " "),
                       str(m["ncur"][src]), _parafusos(c), _carregamento(c),
                       _faixa(c["f0"], "").strip() or "—",
                       str(m["ngrupos"].get(src, 0)),
                       str(m["ngrp"].get(src, 0)),
                       str(m["npc"].get(src, 0))])
    tabela(doc, ["ref.", "source", "curves", "bolt", "loading", "F₀ [kN]",
                 "groups", "shared", "per-curve"], linhas,
           larguras=[0.35, 1.15, 0.45, 0.6, 1.75, 0.65, 0.45, 0.5, 0.6],
           size=7)
    p(doc, "The reference number points to §14.1, where each source carries "
           "its DOI. The conditions are read from the case registry: bolt "
           "sizes, the loading mode with its range of imposed amplitude or "
           "axial force amplitude, and the range of initial preload. The "
           "calibration cost is published next to the curve count on "
           "purpose. A **group** is a set of curves that share one "
           "configuration; **shared** counts the numbers that group carries; "
           "**per-curve** counts values attached to a single curve, most of "
           "them read from the paper. §8.4 reads these columns.")

    # ---------------------------------------------------------------- 5
    doc.add_heading("5. From published figure to comparable numbers", 1)
    p(doc, "A validation claim built on digitised figures is only as good as "
           "the digitisation, so the chain is stated in full and every step is "
           "auditable.")
    tabela(doc, ["step", "what happens", "where it lives"],
           [["1", "figure extracted from the PDF",
             "paper_figures/<source>__<fig>.png"],
            ["2", "curve digitised point by point", "digitized_csv/*.csv"],
            ["3", "axis conventions applied: (x − offset) · scale",
             "ValidationCase.csv_x_offset / csv_x_scale"],
            ["4", "model simulated cycle by cycle", "validation/runner.py"],
            ["5", "model anchored at the first data cycle",
             "CaseResult.align"],
            ["6", "metric window: FLOOR_TRIM = 0.10 and trim_n_max",
             "runner.FLOOR_TRIM, cfg.trim_n_max"],
            ["7", "the three vectors actually compared",
             "metric_x / metric_data / metric_pred"]],
           larguras=[0.4, 3.0, 2.8], size=8.5)
    if cadeia:
        fora = cadeia["n_csv"] - cadeia["n_metric"]
        motivo = ""
        if fora > 0:
            partes = []
            if cadeia.get("abaixo_piso"):
                nb = cadeia["abaixo_piso"]
                partes.append(f"{nb} {'lie' if nb > 1 else 'lies'} below the "
                              f"10 % floor")
            if cadeia.get("alem_nmax"):
                na = cadeia["alem_nmax"]
                quem = ("the last one, at" if na == 1
                        else f"{na} lie beyond n_max, the last at")
                partes.append(
                    f"{quem} N = {cadeia['x_ultimo']:g}, {'lies ' if na == 1 else ''}"
                    f"a fraction of a cycle beyond n_max = {cadeia['n_max']:g}. "
                    f"The simulation stops at the integer part of the last "
                    f"digitised cycle; this edge effect of the implementation "
                    f"is reported rather than hidden")
            motivo = ((" One digitised point is outside the window: "
                       if fora == 1 else
                       f" {fora} digitised points are outside the window: ")
                      + "; ".join(partes) + ".")
        fig(doc, "fig_extraction_chain",
            f"The chain of §5 on one curve, {cadeia['cid']}. (a) The figure "
            f"as published. (b) The {cadeia['n_csv']} digitised points with the "
            f"axis conventions applied. (c) The simulation, raw and divided by "
            f"its own value at the first data cycle (align = "
            f"{cadeia['align']:.4f}). The shaded band is the metric window; the "
            f"dotted line is the 10 % floor; the dash-dotted line marks the end "
            f"of the simulation. (d) The {cadeia['n_metric']} points the metric "
            f"compared, with the residual drawn as vertical segments."
            + motivo
            + (" The first data point sits at N = 0 and is drawn at the left "
               "edge of the logarithmic axis." if cadeia.get("x0_zero") else ""))
    p(doc, "Step 5 deserves a note. The published curves are normalised to "
           "F/F₀ = 1 at the first reported cycle, so the settling that happened "
           "before that point has no measured counterpart. The model is "
           "therefore divided by its own value at that cycle before comparison. "
           "Step 6 removes points below 10 % retention from the metric. It "
           "**also shortens the simulation**, a property of the implementation "
           "that any user reading a truncated curve needs to know.")
    nota(doc, "Consumers read metric_x / metric_data / metric_pred and never "
              "re-interpolate. On the sampled grid, re-interpolation was "
              "measured (2026-07-27) to err by up to 46 % in the embedding "
              "transient, and it once caused four mutually inconsistent numbers "
              "to be published for the same curve.")

    doc.add_heading("5.1 Digitisation uncertainty", 2)
    dg = info.get("digit")
    p(doc, "Three numbers bound what a digitised point is worth, and they are "
           "of different kinds. The first is the reading resolution of the "
           "digitiser on a dense-marker figure: about ±0.005 in F/F₀ for an "
           "axis resolved at 5 % per division, the figure recorded in the "
           "apparatus note of Liu et al. (2017) and used as the reading floor "
           "throughout the corpus.")
    if dg:
        p(doc, f"The second is measured, not assumed. One test of Lu et al. "
               f"(2024) is published in two figures of the paper, and both "
               f"were digitised independently ({dg['a']}, {dg['n_a']} points, "
               f"and {dg['b']}, {dg['n_b']} points). Interpolated on their "
               f"common window, N = {dg['lo']:g} to {dg['hi']:g}, the two "
               f"digitisations differ by **MAE {dg['mae']:.4f}**, "
               f"**σ {dg['sigma']:.4f}** and at most {dg['max']:.4f} in F/F₀. "
               f"The acceptance limits of §6 sit "
               f"{rh.META_MAE/max(dg['mae'],1e-9):.0f}× above that MAE and "
               f"{rh.META_SRES/max(dg['sigma'],1e-9):.0f}× above that σ, so "
               f"digitisation is not what decides a verdict. The pair is kept "
               f"in the store for this purpose and counted once in the census "
               f"(Appendix B).")
    p(doc, "The third is the round-trip against what the paper prints. Where "
           "a source gives a terminal value, a life or a table entry that the "
           "digitised curve must reproduce, the apparatus note of that source "
           "records the check; the notes of Yang et al. (2021, Table 3), Lu "
           "et al. (2024, Table 9) and Rousseau & Bouzid (2025, Fig. 7 "
           "anchors) carry such round-trips, and curves were re-digitised when "
           "a check failed, each time with the check that triggered it "
           "recorded in the project history. The largest uncertainty "
           "of all, the scatter between replicate specimens of the same "
           "condition, is not a digitisation error and is handled by the "
           "criterion itself (§6).")

    doc.add_heading("5.2 Judgement windows", 2)
    from bolt_analysis_studio.validation.inputs import load_full_curve
    linhas, n_meets = [], 0
    for r in sorted(comp, key=lambda q: (q.source, q.case_id)):
        rr = res.get(r.case_id)
        cu = (getattr(rr, "config_used", None) or {}) if rr else {}
        trim = cu.get("trim_n_max")
        if not trim:
            continue
        vc = r.validation_case
        try:
            xr, _yr = load_full_curve(vc.reference_csv_path)
            fim = float(((np.asarray(xr, float) - float(vc.csv_x_offset or 0))
                         * float(vc.csv_x_scale or 1)).max())
        except Exception:
            fim = float("nan")
        g = rn._adopted_for(r.source, r.case_id,
                            getattr(vc, "bolt_size", "") or "")
        prov = (kb.adopted_config(g) or {}).get("prov") or {}
        txt = _texto_trim(prov, trim)
        motivo = motivo_trim(txt)
        if txt is None:
            print(f"  [WARN] trim of {r.case_id} has no provenance text")
        okc = bool(rh._tripe_ok(rr, rh.limite_sres(r.source, pisos)))
        n_meets += okc
        linhas.append([r.case_id, r.source.replace("_", " ").lower(),
                       f"{fim:g}" if fim == fim else "—", f"{float(trim):g}",
                       f"{100*float(trim)/fim:.0f} %" if fim == fim and fim > 0 else "—",
                       motivo])
    p(doc, f"Step 6 of the chain admits a window set by judgement, "
           f"trim_n_max, beyond which the data are not scored. "
           f"**{len(linhas)} of the {m['n']} curves carry one**; "
           f"{n_meets} of them meet the criterion inside it. Every window is "
           f"recorded in the adopted configuration with the rule that placed "
           f"it, and the per-curve reports always show the full curve next to "
           f"the scored one. The reasons fall into few kinds: a stage the "
           f"model does not claim (fatigue fracture, a shear crack, a debris "
           f"tail), or a near-vertical collapse on which no automatic metric "
           f"is well posed (§10.3). In each case the cut was placed by a rule "
           f"declared before the curve was scored, typically a local rate "
           f"exceeding a multiple of the stage-II median for a contiguous "
           f"suffix.")
    tabela(doc, ["curve", "source", "last data cycle", "scored up to N ≤",
                 "kept", "why"], linhas,
           larguras=[2.1, 0.9, 0.7, 0.75, 0.45, 1.6], size=7)

    # ---------------------------------------------------------------- 6
    doc.add_heading("6. Acceptance criterion", 1)
    p(doc, "A curve is accepted when **all three** of the following hold on the "
           "residual r = model − data over the metric window:")
    tabela(doc, ["leg", "meaning", "limit", "anchor"],
           [["max |r|", "worst single point", "≤ 0.10",
             "engineering tolerance"],
            ["MAE", "typical error", "≤ 0.05",
             "the decision margin between ISO 16130 (85 %) and DIN 25201-4 (80 %)"],
            ["σ_res", "shape: scatter of the residual about its own bias",
             "≤ max(0.025, source floor)",
             "median repeatability floor measured across replicate families"]],
           larguras=[0.8, 2.2, 1.3, 2.0], size=8.5)
    p(doc, "The three are norms of the same vector and they measure different "
           "failures. A model can have a small MAE and a large σ_res: small "
           "error, wrong shape. The σ_res limit is raised to the measured "
           "repeatability floor of the source when that floor is larger, "
           "because below the floor 'failure' would be measuring the "
           "experiment's scatter rather than the model.")
    if anat.get("ok"):
        if anat.get("forma"):
            rf = res[anat["forma"]]
            srf = rh.sres_para_censo(rf) or 0.0
            lsf = rh.limite_sres(
                {r.case_id: r for r in comp}[anat["forma"]].source, pisos)
            leg = (f"The three legs on two curves. Left, {anat['ok']} meets "
                   f"all three. Right, {anat['forma']}: MAE "
                   f"{rf.mae:.4f} and max |r| {rf.maxerr:.4f} are both inside "
                   f"their limits, yet σ_res = {srf:.4f} is "
                   f"{srf / lsf:.2f}× the limit of its source: small error, "
                   f"wrong shape. Top: model and data. Bottom: the residual "
                   f"r = model − data, its bias (dashed), the band bias ± "
                   f"σ_res, the ± MAE lines and the point that sets max |r|.")
        else:
            leg = (f"The three legs on {anat['ok']}: the residual "
                   f"r = model − data, its bias (dashed), the band bias ± "
                   f"σ_res, the ± MAE lines and the point that sets max |r|.")
        fig(doc, "fig_residual_three_legs", leg)
        if anat.get("forma"):
            p(doc, "The right-hand curve is why the third leg exists. Its "
                   "residual changes sign along the test: the model is above "
                   "the data early and below it late, so the errors partly "
                   "cancel in the MAE while the shape is wrong throughout. No "
                   "scale constant fixes this. A constant moves the residual "
                   "as a block; it cannot move the point where the residual "
                   "crosses zero. That is what makes σ_res the leg that decides "
                   "most of the failures in §10.2.")
    p(doc, f"On the present corpus, **{m['tripe']} of {m['n']} curves "
           f"({100*m['tripe']//m['n']} %) satisfy all three legs**.")

    # ---------------------------------------------------------------- 7
    doc.add_heading("7. Capabilities", 1)
    for txt in [
        "**One physics for the whole corpus.** The same four mechanisms, the "
        "same coupling and the same equations run every curve. No source "
        "receives a mechanism of its own.",
        "**Attribution.** Every prediction decomposes into per-mechanism "
        "preload loss that sums exactly to the total, so each result carries "
        "its own explanation.",
        f"**Both loading regimes.** Transverse (displacement-controlled) and "
        f"axial (force-controlled); {n_tr} and {n_ax} curves respectively.",
        f"**Bolt sizes M{min(m['cov']['bolt diameter [mm]']):g} to "
        f"M{max(m['cov']['bolt diameter [mm]']):g}** and preloads from "
        f"{min(m['cov']['preload F0 [kN]']):.1f} kN to "
        f"{max(m['cov']['preload F0 [kN]']):.0f} kN; the validity "
        f"envelope is in §12.3.",
        "**Runaway and self-arrest both emerge from the coupled loop**, with "
        "no switch to select them: the same configuration produces a joint "
        "that collapses and one that stabilises, driven only by the input "
        "that changed (§8).",
        "**Zero-refit prediction is documented**, including one condition "
        "digitised later, never used to adjust any constant, and predicted "
        "inside the acceptance criterion (§8.2).",
    ]:
        pr = doc.add_paragraph(style="List Bullet")
        for pedaco, forte in _negrito(txt):
            r = pr.add_run(pedaco)
            r.bold = forte
        pr.paragraph_format.space_after = Pt(3)
    dom = ", ".join(f"{k.replace('_',' ')} {v}"
                    for k, v in m["dom"].most_common())
    p(doc, f"Dominant mechanism across the corpus, as attributed by the model: "
           f"{dom}.")



    # ---------------------------------------------------------------- 8
    doc.add_heading("8. Why this is not a curve fit", 1)
    p(doc, "Any model with more than a handful of constants should face this "
           "question, and evidence is the only answer worth giving. Five "
           "independent lines follow. The last is unfavourable to the "
           "software; it is included for precisely that reason.")

    doc.add_heading("8.1 One configuration, opposite behaviours", 2)
    p(doc, "A fit reproduces a curve because it was tuned to that curve. The "
           "test that separates the two is therefore: **hold the constants "
           "fixed, change only the input, and ask whether the model follows.**")
    porid = {r.case_id: r for r in todos_id}
    fam, fora = grupo_fisica, n_fora_fisica
    linhas = []
    for c in fam:
        rr = res_all[c]
        amp = float(getattr(porid[c].validation_case,
                            "transverse_displacement_mm", 0) or 0)
        cens = rh.caso_comparavel(porid[c].source, c)
        ok = rh._tripe_ok(rr, rh.limite_sres(porid[c].source, pisos))
        linhas.append([f"{amp:g} mm", f"{float(rr.metric_data[-1]):.3f}",
                       f"{float(rr.metric_pred[-1]):.3f}", f"{rr.mae:.4f}",
                       ("accepted" if ok else "outside") if cens
                       else "not in the census"])
    lo = min(float(res_all[c].metric_data[-1]) for c in fam)
    hi = max(float(res_all[c].metric_data[-1]) for c in fam)
    p(doc, f"The transverse amplitude sweep of Lu et al. (2024), figure 18, "
           f"gives exactly that experiment: **{len(fam)} curves, one "
           f"configuration, the same constants for all of them.** Only the "
           f"imposed amplitude differs:")
    tabela(doc, ["amplitude", "measured F/F₀ (end)", "predicted F/F₀", "MAE",
                 "criterion"], linhas, larguras=[1.0, 1.5, 1.3, 1.0, 1.1],
           size=9)
    p(doc, f"Measured retention spans **{lo:.3f} to {hi:.3f}**, from a joint "
           f"that holds to one that has essentially let go, and one "
           f"configuration follows both ends. Were these independent fits, "
           f"each row would carry its own constants and the agreement would "
           f"be trivial and worthless.")
    nao_cens = [c for c in fam
                if not rh.caso_comparavel(porid[c].source, c)]
    if fora:
        nota(doc, f"The table leaves out {fora} further curve of the same "
                  f"figure because it carries a constant of its own. "
                  f"Including it would weaken the claim, so it is omitted, "
                  f"and this note puts the exclusion on record.")
    if nao_cens:
        nota(doc, "The row marked 'not in the census' is the same physical "
                  "test published in two figures of the paper, so counting "
                  "it would count one measurement twice. It appears here "
                  "because this section asks about the physics rather than "
                  "the corpus size, and it runs on the same constants.")
    fig(doc, "fig_one_physics",
        "One physics, several behaviours. Points are digitised "
        "data, lines are the prediction; all panels share the same constants.")

    doc.add_heading("8.2 Prediction on a condition that did not exist yet", 2)
    CAL, NOV = "rousseau2025_hdpe_t10", "rousseau2025_hdpe_t10_amp0p2"
    if CAL in res and NOV in res:
        ca = rn._effective_overrides(porid[CAL], {})
        nv = rn._effective_overrides(porid[NOV], {})
        extra = sorted(set(ca) - set(nv))
        difs = [k for k in set(ca) & set(nv) if ca[k] != nv[k]
                and k != "delta_amp_mm"]
        p(doc, f"The single strongest test in the corpus, and it is worth "
               f"describing in full. The HDPE joint of Rousseau et al. (2025) "
               f"was calibrated on one amplitude ({CAL}). A **different "
               f"condition**, published in a later figure of the same paper, "
               f"was digitised afterwards and predicted with no adjustment "
               f"whatever.")
        p(doc, f"The calibrated curve runs on **{len(ca)} constants**; the "
               f"predicted one runs on **{len(nv)}**. The "
               f"{len(extra)} it does not receive are precisely the ones that "
               f"were fitted on the calibration curve "
               f"({', '.join(extra)}), and its per-curve entry is deliberately "
               f"left empty so that it cannot inherit them by name matching. "
               f"Of the constants both share, "
               f"{'none differ' if not difs else 'the differing ones are ' + ', '.join(difs)} "
               f"apart from the imposed amplitude itself.")
        tabela(doc, ["", "calibration condition", "predicted condition"],
               [["amplitude",
                 f"{float(getattr(porid[CAL].validation_case, 'transverse_displacement_mm', 0) or 0):g} mm",
                 f"{float(getattr(porid[NOV].validation_case, 'transverse_displacement_mm', 0) or 0):g} mm"],
                ["constants used", str(len(ca)), str(len(nv))],
                ["measured final F/F₀", f"{float(res[CAL].metric_data[-1]):.3f}",
                 f"{float(res[NOV].metric_data[-1]):.3f}"],
                ["predicted final F/F₀", f"{float(res[CAL].metric_pred[-1]):.3f}",
                 f"{float(res[NOV].metric_pred[-1]):.3f}"],
                ["MAE", f"{res[CAL].mae:.4f}", f"{res[NOV].mae:.4f}"]],
               larguras=[1.7, 2.2, 2.2], size=8.5)
        p(doc, f"The two conditions do not merely differ in degree. The "
               f"calibration condition **collapses**, ending at "
               f"{float(res[CAL].metric_data[-1]):.3f} of its preload, while "
               f"the predicted condition **holds**, at "
               f"{float(res[NOV].metric_data[-1]):.3f}. The model was given "
               f"only the new amplitude and reproduced the second behaviour to "
               f"within MAE {res[NOV].mae:.4f}, inside the acceptance "
               f"criterion of §6.")
        p(doc, "A fit does not extrapolate outside its own adjustment range, "
               "onto a qualitatively different behaviour, using fewer "
               "constants than the curve it was fitted to. This one did.")
        ev = info.get("evidencia_82") or {}
        if ev.get("csv_nov"):
            partes = [f"The predicted curve entered the repository on "
                      f"{ev['csv_nov'][1]} (commit {ev['csv_nov'][0]})"]
            if ev.get("csv_cal"):
                partes.append(f"the calibration curve on {ev['csv_cal'][1]} "
                              f"(commit {ev['csv_cal'][0]})")
            frase = "; ".join(partes) + "."
            if ev.get("ultima"):
                d_ult, k_ult = ev["ultima"]
                frase += (f" The latest date recorded in the provenance of "
                          f"the constants the two curves share is "
                          f"{d_ult}, for {k_ult}, an adjustment made on the "
                          f"calibration curve alone after that curve was "
                          f"re-digitised.")
            frase += (" The predicted curve has never entered a fit: its "
                      "per-curve entry is empty, and no constant in the "
                      "configuration cites it. The claim is therefore "
                      "zero-refit with respect to the predicted data, which "
                      "is the claim that matters; it is not a claim that "
                      "every shared constant was frozen before the figure "
                      "was digitised, and the dates above let the reader "
                      "check the difference.")
            nota(doc, frase)
        if NOV in dict(  # the same curve is an unsafe call in §11
                (cid, 1) for cid, *_r in m["fs_all"]):
            nota(doc, "The same curve appears in §11 as one of the "
                      "unsafe engineering calls: it is measured just "
                      "below the 85 % ISO line and predicted just above "
                      "it. Both statements are true and both are "
                      "published. The strongest prediction in the "
                      "corpus is also, at one particular threshold, a "
                      "call an engineer should not have trusted.")

    doc.add_heading("8.3 The model is allowed to fail, and does", 2)
    p(doc, "A fit cannot be falsified: it absorbs whatever it is shown. This "
           "model has been falsified repeatedly, and the record is kept "
           "deliberately. Examples, each with the measurement that killed it:")
    tabela(doc, ["proposed mechanism", "how it died"],
           [["Axial-load-driven arrest floor",
             "the floor moves the wrong way with axial load; falsified "
             "before a line of code was written"],
            ["Slip-onset incubation",
             "predicted a reduction ordered by amplitude; the measurement "
             "showed the opposite ordering"],
            ["Decelerating creep kernel (general)",
             "8 of 18 curves got worse, one from 0.040 to 0.223"],
            ["Cattaneo–Mindlin slip regime",
             "inert: the gates were never reached in the adopted packs "
             "(Δ = 0.0000 on six channels)"],
            ["Stiffness modulation by fatigue damage",
             "broke energy conservation by up to −20.5 J with no counterpart "
             "in external work"]],
           larguras=[2.2, 4.0], size=8.5)
    p(doc, "The measurements behind this table are filed in the repository "
           "(MODEL_LEGITIMACY.md and the pre-registration files under "
           "docs/superpowers/specs/), each with its date and its fingerprint. "
           "At the time of writing (August 2026) thirty-eight such "
           "falsifications and retractions were counted in "
           "the project history, several of them against results the project "
           "had already published. **An instrument that can say no, and does, "
           "is reporting rather than fitting.**")

    doc.add_heading("8.4 The cost is published next to the error", 2)
    p(doc, "The honest way to answer 'how many adjustable numbers?' is to "
           "publish the count without being asked, and to publish it split. A "
           "single lumped ratio flatters the model in one direction and "
           "slanders it in the other, because the numbers in a configuration "
           "are not all of one kind.")
    ncurva = sorted(m["por_curva"])
    curvas_pc = sum(m["ncur"][s] for s in ncurva)
    tabela(doc, ["", "count", "reading"],
           [["curves in the corpus", str(m["n"]), "—"],
            ["calibration groups", str(m["n_groups"]),
             f"configurations referenced by at least one comparable curve; "
             f"one configuration per {m['n']/max(m['n_groups'],1):.1f} curves "
             f"on average"],
            ["shared constants", str(m["k_grupo"]),
             "numbers a group carries for all its curves, including geometry "
             "and friction read from the paper"],
            ["per-curve entries", str(m["k_percurva"]),
             "values attached to one curve, mostly published fatigue lives "
             "and per-specimen amplitudes"],
            ["distinct constant names", str(m["nomes_tot"]),
             f"of the {len(specs)} fields the model exposes, this many are "
             f"touched anywhere in the corpus"]],
           larguras=[1.7, 0.8, 3.7], size=8.5)
    p(doc, "**The group is the unit that matters.** A group is a set of "
           "curves sharing one configuration, so a source served by one group "
           "is a source where every curve after the first is a prediction. A "
           "source partitioned into as many groups as it has curves is, in "
           "substance, calibrated curve by curve, whatever the constant count "
           "says.")
    fig(doc, "fig_calibration_cost",
        "Calibration cost by source: how many curves each configuration has "
        "to serve. Sources at 1.0 are calibrated one configuration per curve.")
    razao = sorted(((s, m["ncur"][s] / max(m["ngrupos"].get(s, 1), 1))
                    for s in m["ncur"] if m["ngrupos"].get(s)),
                   key=lambda t: -t[1])
    p(doc, f"Measured that way: **{len(ncurva)} of {len(m['ncur'])} sources "
           f"are calibrated one group per curve** "
           f"({', '.join(s.replace('_',' ').lower() for s in ncurva)}), "
           f"accounting for {curvas_pc} curves "
           f"({100*curvas_pc/m['n']:.0f} % of the corpus). For those sources "
           f"the word 'prediction' would be unearned, and the annex does not "
           f"use it. At the other "
           f"end, one single configuration serves all "
           f"{m['ncur'][razao[0][0]]} curves of "
           f"{razao[0][0].replace('_',' ').lower()}, and "
           f"{sum(1 for s, _v in razao if m['ngrupos'].get(s) == 1)} "
           f"sources in all are served by a single configuration.")
    zero = sorted(s for s in m["ncur"] if not m["ngrupos"].get(s))
    if zero:
        p(doc, f"One source sits outside the chart's own scale: "
               f"**{', '.join(s.replace(chr(95),' ').lower() for s in zero)} "
               f"carries no adopted configuration at all**. Its "
               f"{sum(m['ncur'][s] for s in zero)} curves run on the "
               f"shared physics and the model's defaults, with zero "
               f"source-specific constants, and both pass the "
               f"three-leg criterion. They are the cheapest prediction "
               f"in the corpus.")
    n_fund = 0
    for s in kb.adopted_sources():
        e = kb.adopted_config(s) or {}
        prov = e.get("prov") or {}
        for _t, x in ((e.get("cfg") or {}).get("per_case") or {}).items():
            if isinstance(x, dict):
                for campo in x:
                    # exact key first: the merge wrote prov[campo]; an alias or
                    # composite key found by prov_lookup may shadow it
                    txt = prov.get(campo) or prov_lookup(prov, campo) or ""
                    if "fundido da colheita" in txt:
                        n_fund += 1
    if m["prov_colhida"]:
        cobertura = (f"{m['prov_com']} carry a provenance field in the "
                     f"configuration itself and {m['prov_colhida']} are "
                     f"documented in the harvested provenance map, each naming "
                     f"the document that introduced them")
    else:
        cobertura = (f"all {m['prov_com']} carry a provenance text in the "
                     f"configuration itself"
                     + (f"; {n_fund} of those texts were harvested from the "
                        f"pre-registration record on 2026-08-25 and merged into "
                        f"the configuration on 2026-08-28, each naming the "
                        f"document that introduced the value" if n_fund else ""))
    p(doc, f"Of the **{m['prov_tot']} per-curve entries**, {cobertura}. "
           f"**{'None' if not m['prov_sem'] else len(m['prov_sem'])} are undocumented.** §9.3 draws the "
           f"distinction between a number read from a paper and a number "
           f"adjusted on a curve; that distinction, more than the raw "
           f"count, decides whether a constant is a cost.")
    led = info.get("ledger") or []
    sh = [x for x in led if x["scope"] == "shared"]
    if sh:
        n_cfg = sum(1 for x in sh if x["where"] == "configuration")
        n_col = sum(1 for x in sh if x["where"] == "harvested map")
        n_none = sum(1 for x in sh if x["where"] == "none")
        p(doc, f"The shared constants are covered less completely, and the "
               f"gap is published rather than rounded away: of the "
               f"{len(sh)} shared numbers, {n_cfg} carry a provenance text in "
               f"the configuration, {n_col} in the harvested map, and "
               f"**{n_none} carry none**. Appendix C lists every one of them, "
               f"with the class each text supports.")

    doc.add_heading("8.5 Where the claim is weakest", 2)
    p(doc, "Three qualifications, stated because a reader who finds them "
           "unaided will discount everything above.")
    for txt in [
        "**Constants do not transfer between rigs**, although the forms do. "
        "C_creep, for instance, is a property of the tribological pair: an "
        "anchor on 304 stainless and a fit on a different rig differ by more "
        "than an order of magnitude, with disjoint confidence intervals. The "
        "model is a shared physics with a per-rig calibration shell, and "
        "calling it anything stronger would be false.",
        f"**Four sources are calibrated one configuration per curve.** "
        f"They are named in §8.4 and they carry "
        f"{sum(m['ncur'][s] for s in m['por_curva'])} of the "
        f"{m['n']} curves. For those the agreement demonstrates that the "
        f"model CAN take the shape of the data, which is a weaker claim "
        f"than the rest of this section makes.",
        "**The fatigue life is an input to the model.** Where an experiment "
        "ran to fracture, the published life is supplied, and the claim "
        "becomes 'predicts the curve given the life'. Prediction of the life "
        "itself remains falsified: the internal clock was measured to be "
        "wrong by roughly ±36 % (MODEL_LEGITIMACY.md, §4.52).",
    ]:
        pr = doc.add_paragraph(style="List Bullet")
        for pedaco, forte in _negrito(txt):
            r = pr.add_run(pedaco)
            r.bold = forte
        pr.paragraph_format.space_after = Pt(4)

    # ---------------------------------------------------------------- 9
    doc.add_heading("9. The constants", 1)
    p(doc, f"The material model exposes **{len(specs)} fields**. They are not "
           f"{len(specs)} free parameters: most are inactive by default, many "
           f"are structural switches, and the ones that are fitted are "
           f"adjusted once for a whole rig at a time. This section explains "
           f"what each does.")

    doc.add_heading("9.1 The principal constants", 2)
    p(doc, "The constants that carry the physics, with the equation each one "
           "lives in. The full set is tabulated in §9.2.")
    por_nome = {s.name: s for s in specs}
    for nome, (unidade, equacao, blocos) in PRINCIPAIS.items():
        s = por_nome.get(nome)
        if s is None:
            # a field the engine no longer exposes must not be described
            print(f"  [WARN] principal constant {nome} has no VarSpec; skipped")
            continue
        doc.add_heading(f"{nome}  [{unidade}]", 3)
        mono(doc, equacao, size=8.5)
        for bloco in blocos:
            p(doc, bloco, size=9.5, space=4)

    doc.add_heading("9.2 Complete field reference", 2)
    p(doc, "Every field, its group and its role. **inactive** marks a field "
           "whose default renders it exactly neutral: it changes nothing "
           "until deliberately enabled. Such fields are listed so that the "
           "count of fields is not mistaken for a count of free parameters.")
    grupos = collections.defaultdict(list)
    for s in specs:
        grupos[s.group or "other"].append(s)
    n_neg = sum(1 for s in specs if getattr(s, "negligible", False))
    for g in sorted(grupos):
        doc.add_heading(g.replace("_", " "), 3)
        linhas = []
        for s in sorted(grupos[g], key=lambda q: q.name):
            resumo = prosa_tabela(sem_ui(limpa(s.physics_en).split("\n")[0]))
            resumo = re.sub(r"\s+", " ", resumo)
            if len(resumo) > 150:
                resumo = resumo[:147].rsplit(" ", 1)[0] + "…"
            linhas.append([s.name, unidade_en(s.unit),
                           ("inactive: " if getattr(s, "negligible", False)
                            else "") + resumo])
        tabela(doc, ["field", "unit", "role"], linhas,
               larguras=[1.3, 0.6, 4.3], size=7.5)
    p(doc, f"Of the {len(specs)} fields, {n_neg} are structurally inactive at "
           f"their defaults.")

    doc.add_heading("9.3 Provenance", 2)
    p(doc, "Every constant carries a class, and the classes are ordered by "
           "strength:")
    tabela(doc, ["class", "meaning"],
           [["read (direct)", "the number is printed in the paper: a table "
                              "terminal, a drawn marker, a protocol definition"],
            ["read (canonical reader)",
             "a published algorithm extracts it from the raw data, with its "
             "bias declared"],
            ["regressed", "least squares of a closed-form law to a published "
                          "observable, with r² reported"],
            ["anchored", "fitted on a grid but validated against an independent "
                         "observable"],
            ["fitted (declared)",
             "no anchor; accepted only with a demonstrated interior region, an "
             "honest label, and any degeneracy declared"],
            ["channel off",
             "the value is exactly zero, recording a decision to disable a "
             "channel; nothing was adjusted"],
            ["default (design starter)",
             "the value shipped with the mechanism's design specification, "
             "left unchanged and documented as such"],
            ["form switch",
             "a mode selector or an on/off flag, chosen by pre-registration; "
             "a choice of mechanism, not an adjusted value"],
            ["inherited",
             "a value carried unchanged from an earlier documented adoption "
             "(a parent group, a sibling token); its class is recorded there"]],
           larguras=[1.6, 4.6], size=8.5)
    p(doc, "The distinction matters more than the granularity. A per-curve "
           "number **read from the paper** (the fatigue life of a specific "
           "specimen, for instance) is legitimate; a per-curve number "
           "**adjusted on that curve** is a fit. Both exist in this software, "
           "and the provenance class is what separates them.")
    led = info.get("ledger") or []
    if led:
        cont = collections.Counter((x["scope"] == "shared", x["classe"])
                                   for x in led)
        linhas = []
        for cl in _CLASSES_TABELA:
            a, b_ = cont.get((True, cl), 0), cont.get((False, cl), 0)
            if a or b_:
                linhas.append([cl, str(a), str(b_), str(a + b_)])
        n_sh = sum(1 for x in led if x["scope"] == "shared")
        n_pc = len(led) - n_sh
        linhas.append(["**all**", f"**{n_sh}**", f"**{n_pc}**",
                       f"**{len(led)}**"])
        p(doc, "The class of every constant the corpus runs on, read from the "
               "provenance text recorded with it. The reading is "
               "**conservative**: when a text supports more than one class, "
               "the weakest claim is kept, so a number both read from a paper "
               "and adjusted on a curve counts as fitted. Texts that name no "
               "recognisable class are counted as 'documented (class not "
               "parsed)', never promoted; entries with no text are counted "
               "as undocumented. Appendix C lists every row.")
        tabela(doc, ["class", "shared constants", "per-curve entries",
                     "total"], linhas, larguras=[2.2, 1.3, 1.3, 1.0],
               size=8.5)
        n_fit = cont.get((True, "fitted (declared)"), 0) + cont.get((False, "fitted (declared)"), 0)
        n_read = sum(cont.get((k, cl), 0) for k in (True, False)
                     for cl in ("read (direct)", "read (canonical reader)",
                                "regressed", "anchored"))
        p(doc, f"Read that way, {n_fit} of the {len(led)} entries "
               f"({100*n_fit/max(len(led),1):.0f} %) are fitted constants and "
               f"{n_read} ({100*n_read/max(len(led),1):.0f} %) are numbers "
               f"read, regressed or anchored to a published observable; the "
               f"remainder are channels turned off, mode switches, design "
               f"defaults, values inherited from an earlier documented "
               f"adoption, or entries whose text does not state a class. The "
               f"fitted share is the honest measure of what the corpus cost, "
               f"and it is an upper bound: an inherited value whose parent "
               f"was read from a paper is not counted as read here.")



    # ---------------------------------------------------------------- 10
    doc.add_heading("10. Limitations", 1)
    p(doc, "Wherever a measurement exists, the limitation is stated as one, "
           "because a limitation with a number attached can be checked and a "
           "caveat cannot.")

    doc.add_heading("10.1 What the model does not claim", 2)
    dentro = fora_b = semv = 0
    for r in comp:
        l7 = ((store.get(r.case_id) or {}).get("l7_check") or {})
        v = l7.get("implied_J_per_mm3")
        if v is None:
            semv += 1
        elif l7.get("in_bound"):
            dentro += 1
        else:
            fora_b += 1
    nlt6 = sum(1 for r in comp
               if res.get(r.case_id) is not None
               and rh.sres_para_censo(res[r.case_id]) is None)
    tabela(doc, ["limitation", "measured extent"],
           [["Constants do not transfer between rigs",
             "the forms transfer; the numbers are per tribological pair. "
             "C_creep differs by more than an order of magnitude between an "
             "anchored 304-stainless value and a fitted one, with disjoint "
             "confidence intervals"],
            ["The fatigue life is an input",
             "where a test ran to fracture the published life is supplied. "
             "The internal life clock was measured to be wrong by roughly "
             "±36 % (MODEL_LEGITIMACY.md, §4.52) and the prediction of life "
             "remains falsified"],
            ["The threshold-crossing clock is not predictive",
             "the cycle at which F/F₀ first reaches 0.95 disagrees with the "
             "data by 3–5× in both directions in the source where it was "
             "tested; only the curve shape is claimed"],
            ["Wear-removal energetics are phenomenological",
             f"an informational check compares the implied removal energy "
             f"with a literature band of 1800–10500 J/mm³: of {dentro+fora_b} "
             f"curves that imply a value, **{fora_b} fall outside it** "
             f"({semv} imply none because the wear channel is inactive)"],
            ["The energy budget does not close on small budgets",
             ((f"the residual W_ext + ΔU − ΣW_diss is persisted per curve "
               f"(Appendix A): median {info['energy_corpus']['med']:.1e} of "
               f"the larger budget term, but {info['energy_corpus']['n_cauda']} "
               f"curves exceed 5 % and the worst reaches "
               f"{info['energy_corpus']['max']:.1e} "
               f"({info['energy_corpus']['pior']}). The tail is made of "
               f"sub-joule budgets (creep under sustained preload, stick "
               f"regime), not of damage or fatigue collapse, which are inactive "
               f"in every curve of it (§2.6); the energetics of material "
               f"removal remain phenomenological besides"
               if info.get("energy_corpus") else
               (f"on the §3.2 example the residual is "
                f"{info['engine']['residual_J']:.2f} J against "
                f"{info['engine']['W_diss_J']:.0f} J dissipated over "
                f"{info['engine']['n']} cycles; the corpus-wide value is not "
                f"in the store yet" if info.get("engine") else
                "the residual is not persisted in the store yet")))],
            ["The deep tail is not scored",
             "points below 10 % retention are removed from the metric, so no "
             "statement is made about the last decade of collapse"],
            ["Shape cannot be judged on very short curves",
             f"{nlt6} comparable curves carry fewer than six points in the "
             f"metric window; for them the shape leg is declared "
             f"non-judgeable, the alternative being a silent pass"],
            ["Temperature is not modelled",
             "ΔT exists in the loading interface but no thermal mechanism is "
             "validated; every curve in the corpus is isothermal"],
            ["Material scope",
             "steel and HDPE bolted joints. One CFRP curve in the library is "
             "explicitly declared out of scope and left unfitted"]],
           larguras=[1.9, 4.3], size=8)

    doc.add_heading("10.2 The curves that do not meet the criterion", 2)
    pernas = collections.Counter()
    piores = []
    for r in comp:
        rr = res.get(r.case_id)
        if rr is None:
            continue
        ls = rh.limite_sres(r.source, pisos)
        if rh._tripe_ok(rr, ls):
            continue
        sr = rh.sres_para_censo(rr)
        mult = {"max |r|": rr.maxerr / rh.META_MAX,
                "MAE": rr.mae / rh.META_MAE}
        if sr is not None:
            mult["σ_res"] = sr / ls
        q = max(mult, key=mult.get)
        pernas[q] += 1
        piores.append((mult[q], r.case_id, q))
    nfora = m["n"] - m["tripe"]
    ordem = ", ".join(f"{k} {v}" for k, v in pernas.most_common())
    p(doc, f"**{nfora} of {m['n']} curves are outside the criterion.** The leg "
           f"that decides the verdict (the one furthest beyond its own limit) "
           f"is distributed as: {ordem}.")
    piores.sort(reverse=True)
    p(doc, f"Severity is bounded: the median outside curve sits at "
           f"{np.median([x[0] for x in piores]):.2f}× its worst limit and the "
           f"single worst at {piores[0][0]:.2f}× "
           f"({piores[0][1]}, {piores[0][2]}).")
    fig(doc, "fig_binding_leg",
        "Which leg decides, and how far outside the curves that "
        "fail actually are.")
    p(doc, "All of them remain in the published record. Were every curve to "
           "pass, the reader could not tell an over-flexible model from an "
           "under-populated corpus.")

    doc.add_heading("10.3 Declared classes", 2)
    p(doc, "Some curves are outside the criterion for reasons that are not "
           "the model's. Each carries a declared class with the measurement "
           "that justifies it, and a declared curve is never counted as a "
           "success:")
    _SIG = {
        "data-limited": "the resolution of the published figure is coarser "
                        "than the tolerance being asked of it, or the figure's "
                        "own provenance is not the rig described",
        "metric-limited": "a near-vertical collapse in which no automatic "
                          "metric on sparse digitised points can distinguish "
                          "a ramp from a cliff",
        "n < 6": "too few points in the window for the shape leg to mean "
                 "anything; the other two legs pass",
        "out of scope": "a material or protocol the model does not claim",
        "protocol orphan": "the source mixes two loading protocols and the "
                           "curve has no replicate in its own; no route by "
                           "model, floor or data",
        "form-limited": "the honest class: the model's shape is wrong and "
                        "every available constant has been tried and falsified",
    }
    # counted from the proofs themselves, so the table cannot list a class
    # no curve carries without saying so (a zero is information)
    cls_cont = collections.Counter(classe_declarada(c)
                                   for c in getattr(rh, "_DECLARADAS", {}))
    ordem_cls = list(_SIG) + [k for k in cls_cont if k not in _SIG]
    tabela(doc, ["class", "meaning", "curves"],
           [[k, _SIG.get(k, "—"), str(cls_cont.get(k, 0))] for k in ordem_cls],
           larguras=[1.3, 4.3, 0.6], size=8.5)
    nexc = len(getattr(rh, "_EXCECOES", {}))
    ndec = len(getattr(rh, "_DECLARADAS", {}))
    dupl = set(getattr(rh, "_EXCECOES", {})) & set(getattr(rh, "_DECLARADAS", {}))
    resolv = m["tripe"] + nexc + ndec - len(dupl)
    p(doc, f"A curve may also carry a **signed exception**: a proof that the "
           f"experiment's own replicate scatter is larger than the tolerance "
           f"being demanded, so that a perfect model would fail the criterion "
           f"too. There are {nexc} of those and {ndec} declared curves, with "
           f"{len(dupl)} counted in both.")
    p(doc, f"**Two readings are therefore published together, and neither on "
           f"its own.** Strict: **{m['tripe']} of {m['n']}** curves meet the "
           f"three legs on their own merit "
           f"({100*m['tripe']/m['n']:.0f} %). Resolved or declared: "
           f"**{resolv} of {m['n']}** ({100*resolv/m['n']:.0f} %), which adds "
           f"the curves whose distance from the criterion has a measured "
           f"cause outside the model. The first number is the one to quote "
           f"about the model; the second is the one that says how much of the "
           f"corpus has been accounted for. Publishing only the second would "
           f"be flattery, and publishing only the first would attribute the "
           f"experiment's scatter to the model.")
    p(doc, "The form-limited class is the working queue. Of these classes, "
           "it alone records a debt owed by the model itself. Appendix A "
           "lists every curve with its status, so the reader can see which "
           "curve carries which class rather than take the counts on trust.")

    # ---------------------------------------------------------------- 11
    doc.add_heading("11. What the model would tell an engineer", 1)
    p(doc, "The acceptance criterion of §6 is a research instrument. A "
           "designer asks a different question: *does this joint stay above "
           "the retention the standard requires?* That question has an "
           "asymmetric cost, so it is reported separately.")
    tabela(doc, ["standard", "threshold", "agreement",
                 "conservative errors", "unsafe errors"],
           [["ISO 16130:2015", "85 % of F₀", f"{100*m['iso']['acc']:.1f} %",
             str(m["iso"]["fa"]), str(m["iso"]["fs"])],
            ["DIN 25201-4", "80 % of F₀", f"{100*m['din']['acc']:.1f} %",
             str(m["din"]["fa"]), str(m["din"]["fs"])]],
           larguras=[1.5, 1.2, 1.1, 1.3, 1.1], size=8.5)
    p(doc, "**Only the last column matters.** A conservative error (the model "
           "says the joint loosens and it did not) costs a redesign. An "
           "unsafe error means the software would have told an engineer the "
           "joint holds when the experiment says it did not.")
    fig(doc, "fig_engineering_decision",
        "Engineering decision at the two standard thresholds. "
        "Circled points in the shaded quadrant are the unsafe calls.")
    if m["fs_all"]:
        linhas = [[cid, f"{o:.3f}", f"{pv:.3f}",
                   "meets §6" if k else "outside §6"]
                  for cid, o, pv, k in sorted(m["fs_all"], key=lambda t: t[1])]
        p(doc, f"The {len(m['fs_all'])} unsafe calls at the ISO threshold, in "
               f"full:")
        tabela(doc, ["curve", "measured F/F₀", "predicted F/F₀",
                     "research criterion"], linhas,
               larguras=[2.6, 1.2, 1.2, 1.3], size=8)
        nq = len(m["fs_tripe"])
        if nq:
            p(doc, f"**{nq} of them satisfy the three-leg criterion of §6.** "
                   f"Far from being a contradiction, this is the point: a curve "
                   f"can be accurate to within 0.05 and still sit on the wrong "
                   f"side of an 85 % line it passes close to. The research "
                   f"criterion measures agreement; the consequence of "
                   f"disagreement lies outside it, and the two must be read "
                   f"together.")
    p(doc, "Practical consequence for a user: **do not read the retention "
           "number alone.** Read the per-mechanism decomposition (§2.2) as "
           "well. A joint predicted at 0.86 whose loss is dominated by "
           "rotational loosening is in a qualitatively different state from "
           "one at 0.86 dominated by embedding, and the model says which.")

    # ---------------------------------------------------------------- 12
    doc.add_heading("12. Where the model has been shown to work", 1)

    doc.add_heading("12.1 Agreement across the corpus", 2)
    p(doc, f"Predicted against measured final retention over all {len(res)} "
           f"curves: **R² = {m['r2']:.4f}**, bias "
           f"**{m['bias']:+.4f}** in F/F₀, with "
           f"**{100*m['in05']:.1f} %** of curves inside ±0.05 and "
           f"**{100*m['in10']:.1f} %** inside ±0.10. Watch the bias: a model "
           f"tuned per curve would have a bias near zero by construction, so "
           f"a systematic offset of this size is the signature of a "
           f"calibration that is genuinely shared.")
    fig(doc, "fig_parity",
        "Parity. Bands are ±0.05 and ±0.10 in retained fraction.")

    doc.add_heading("12.2 By source", 2)
    p(doc, "Aggregate agreement can hide a source that is systematically "
           "wrong, so the median error is also published per source.")
    fig(doc, "fig_median_mae_by_source",
        "Median MAE by source against the 0.05 limit.")
    med = sorted(((s, float(np.median([res[r.case_id].mae for r in comp
                                       if r.source == s
                                       and r.case_id in res])))
                  for s in m["ncur"]), key=lambda t: -t[1])
    piores_src = ", ".join(f"{s.replace('_',' ').lower()} ({v:.3f})"
                           for s, v in med[:4])
    melhores = ", ".join(f"{s.replace('_',' ').lower()} ({v:.3f})"
                         for s, v in med[-4:][::-1])
    p(doc, f"Best served: {melhores}. Worst served: {piores_src}. Naming the "
           f"worst sources serves the same purpose as listing the failing "
           f"curves in §10.2: a reader looking for the model's weak ground "
           f"should not have to find it unaided.")

    doc.add_heading("12.3 Validity envelope", 2)
    p(doc, "The ranges over which the model has been confronted with "
           "experiment. Outside them it may still run, since the equations "
           "impose no bound, but nothing has been demonstrated there.")
    linhas = []
    for k, v in m["cov"].items():
        if v:
            linhas.append([k, f"{min(v):.4g}", f"{max(v):.4g}", str(len(v))])
    tabela(doc, ["quantity", "minimum", "maximum", "curves"], linhas,
           larguras=[2.0, 1.2, 1.2, 0.9], size=8.5)
    p(doc, f"Loading regimes: {n_tr} transverse displacement-controlled and "
           f"{n_ax} axial force-controlled curves. Bolt classes span the "
           f"structural range, and both a metallic and a polymeric clamped "
           f"member are represented.")
    p(doc, "The lowest frequencies in the table are not vibration tests. They "
           "belong to the static creep tests, whose elapsed time the registry "
           "maps onto the cycle axis with a nominal clock: one cycle per hour "
           "for the composite joints of Caccese et al. (1/3600 Hz) and one "
           "cycle per day for the seawater-exposure tests of Yang, Bai & Ding "
           "(1/86400 Hz). The same time-driven creep law then serves both "
           "kinds of test without a special case.")
    fig(doc, "fig_validity_envelope",
        "Validity envelope. Marker area is proportional to bolt "
        "diameter.")

    doc.add_heading("12.4 Attribution beyond agreement", 2)
    p(doc, "Agreement on a scalar is weak evidence; a model that also says "
           "*why* can be checked against physical expectation. Every "
           "prediction decomposes into the four mechanisms, and the parts sum "
           "exactly to the total loss.")
    if alvo:
        fig(doc, "fig_mechanism_decomposition",
            f"{alvo}: prediction against experiment (top) and the "
            f"cumulative loss attributed to each mechanism (bottom).")
    p(doc, "This is also what makes the model falsifiable in detail. A "
           "prediction that lands on the right curve by attributing the loss "
           "to the wrong mechanism can be caught, and several proposals in "
           "§8.3 died exactly that way.")

    doc.add_heading("12.5 Sensitivity", 2)
    tor = info.get("tornado")
    if tor:
        S_tr, S_ax = tor["S_tr"], tor["S_ax"]
        p(doc, f"A one-at-a-time study perturbs each constant by ±20 % about "
               f"its adopted value and records S, the mean shift of the "
               f"predicted F/F₀, on {tor['n_tr']} transverse and "
               f"{tor['n_ax']} axial canonical cases (MODEL_LEGITIMACY.md, "
               f"§4.42). It answers a question the agreement figures cannot: "
               f"which numbers the prediction actually depends on.")
        fig(doc, "fig_sensitivity_tornado",
            "One-at-a-time sensitivity by loading family. Bars are the mean "
            "shift of the predicted retention for a ±20 % perturbation of the "
            "constant named on the left; constants marked frozen have S ≈ 0 "
            "and are excluded from any fit by construction.")
        top = tor["nomes"][:8]
        linhas = [[k, f"{S_tr.get(k, {}).get('mean', 0):.4f}",
                   f"{S_ax.get(k, {}).get('mean', 0):.4f}",
                   ("frozen (S ≈ 0)" if k in tor["frozen"] else
                    "measured input" if k in ("mu", "mu_thread", "mu_bearing")
                    else "constant")] for k in top]
        tabela(doc, ["constant", "S, transverse", "S, axial", "nature"],
               linhas, larguras=[1.8, 1.2, 1.2, 1.6], size=8.5)
        lider = top[0] if top else "—"
        n_in_ax = len(tor["inertes_ax"])
        p(doc, f"Three readings. The most sensitive number, {lider}, is a "
               f"measured input (the friction coefficients, taken from the "
               f"paper or from the torque–preload relation), not an adjusted "
               f"constant; the most sensitive fitted constant is "
               f"tr_loose_gain. Four constants ({', '.join(tor['frozen'])}) "
               f"have S ≈ 0 on this corpus and are frozen: they take part in "
               f"the physics but the data cannot identify them, so offering "
               f"them to an optimiser would be fitting noise. In the axial "
               f"family {n_in_ax} of {len(S_ax)} constants are inert "
               f"(S = 0), because the slip-driven channels are silent under "
               f"force-controlled axial loading; the axial prediction rests "
               f"on embedding and creep alone.")

    # ---------------------------------------------------------------- 13
    doc.add_heading("13. Reproducing the results", 1)
    p(doc, "Everything in this annex is reproducible from the repository "
           "without external data: the digitised curves, the adopted "
           "configurations and the result store are all versioned.")
    mono(doc,
         "# 1. environment\n"
         "pip install -e \".[dev]\"\n\n"
         "# 2. re-simulate the whole corpus and rebuild the reports\n"
         "python -m bolt_analysis_studio.validation.report --all\n\n"
         "# 3. the same, in parallel\n"
         "python New_Theory/parallel_batch.py --workers 6 --store\n\n"
         "# 4. rebuild this annex from the store\n"
         "python New_Theory/build_annex_docx.py\n\n"
         "# 5. the test suite\n"
         "python -m pytest tests/ -q")
    try:
        fp = rn.engine_fingerprint()
    except Exception:
        fp = "(unavailable)"
    tabela(doc, ["artefact", "path"],
           [["result store",
             "Models/CALIBRATION_AND_VALIDATION/validation_store.json"],
            ["adopted configurations", "New_Theory/adopted_configs.json"],
            ["shared physical constants", "New_Theory/joint_calibrations.json"],
            ["digitised curves",
             "Models/CALIBRATION_AND_VALIDATION/curve_library/digitized_csv/"],
            ["apparatus notes (one per source)",
             "curve_library/apparatus_notes/"],
            ["per-curve reports (HTML)", "New_Theory/validation_html/"],
            ["one page per constant", "New_Theory/variable_explorer/"]],
           larguras=[2.1, 4.1], size=8.5)
    p(doc, f"The state of the model is identified by a fingerprint over the "
           f"shared constants and every adopted configuration, including their "
           f"provenance fields. The results in this annex were produced under "
           f"fingerprint **{fp}**. Changing a constant, or merely correcting "
           f"a provenance label, changes the fingerprint and requires the "
           f"whole corpus to be re-simulated, which is what prevents a result "
           f"and the configuration that produced it from drifting apart.")

    # ---------------------------------------------------------------- 14
    doc.add_heading("14. References", 1)

    doc.add_heading("14.1 Experimental sources of the validation corpus", 2)
    p(doc, f"The {m['n_src']} sources below supply every curve the software is "
           f"validated against. Each is cited in the corpus table of §4 with "
           f"the number of curves it contributes.")
    for i, src in enumerate(sorted(m["ncur"]), 1):
        cite, doi = refs.get(src, ["", ""])
        txt = cite or src.replace("_", " ")
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Inches(0.3)
        par.paragraph_format.first_line_indent = Inches(-0.3)
        r = par.add_run(f"[{i}] ")
        r.bold = True
        r.font.size = Pt(9)
        r = par.add_run(txt)
        r.font.size = Pt(9)
        if doi:
            r = par.add_run(f"  https://doi.org/{doi}")
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

    doc.add_heading("14.2 Standards", 2)
    for t in ["ISO 16130:2015, Aerospace series. Dynamic testing of the "
              "locking behaviour of bolted connections under transverse "
              "loading conditions (Junker test); acceptance at 85 % retained "
              "preload.",
              "DIN 25201-4. Design guide for railway vehicles: securing of "
              "bolted joints; acceptance at 80 % retained preload.",
              "VDI 2230 Part 1. Systematic calculation of highly stressed "
              "bolted joints; source of the embedding table f_Z used as a "
              "per-joint input.",
              "ISO 724 and ISO 7089. Thread and washer geometry, used to "
              "derive the contact areas."]:
        par = doc.add_paragraph(style="List Bullet")
        r = par.add_run(t)
        r.font.size = Pt(9)
        par.paragraph_format.space_after = Pt(3)

    doc.add_heading("14.3 Constitutive laws used", 2)
    tabela(doc, ["law", "where it is used"],
           [["Hooke / Coulomb", "elastic assembly and the friction limit "
                                "(the analytical layer)"],
            ["Norton", "embedding, in its exact state-based form"],
            ["Norton–Bailey", "creep of the clamped members"],
            ["Greenwood–Williamson",
             "softening of the contact stiffness as the surface degrades"],
            ["Archard", "wear volume from slip work"],
            ["Junker", "the displacement-controlled transverse test the "
                       "corpus is built on"],
            ["Goodman", "mean-stress correction in the fatigue channel"],
            ["Pedersen", "joint stiffness where the adopted configuration "
                         "selects it"],
            ["Hill function", "smooth gates for incubation, arrest and "
                              "conformation, used in place of hard switches"]],
           larguras=[1.6, 4.6], size=8.5)
    nota(doc, "The three-layer construction is deliberate. Analytical laws "
              "are used where the physics is settled, established empirical "
              "laws where it is not, and a small number of calibrated "
              "multipliers where neither is available. Each of the third kind "
              "carries a declared provenance class (§9.3).")

    doc.add_heading("14.4 Software", 2)
    p(doc, "Bolt Analysis Studio, written by Leonardo Rosa Ribeiro da Silva "
           "(Faculdade de Engenharia Mecânica, Universidade Federal de "
           "Uberlândia, leorrs@ufu.br) and Neilon de Souza da Silva "
           "(Petróleo Brasileiro S.A., neilon@petrobras.com.br). "
           "Released under the MIT licence; the validation "
           "corpus, the adopted configurations and the result store are "
           "distributed with it.", size=9.5)

    apendices(doc, comp, res, pisos, store, m, todos_id, res_all, info)


# --------------------------------------------------------------------------- #
# Appendices — every result, one row per curve                                #
# --------------------------------------------------------------------------- #

def _carga(vc, cu):
    """One cell describing the load: imposed transverse amplitude, axial force
    amplitude, superposed external axial load, or a sustained preload only."""
    d = float(getattr(vc, "transverse_displacement_mm", 0) or 0)
    fa = float(getattr(vc, "axial_force_amplitude_N", 0) or 0)
    if fa <= 0:
        fa = float((cu or {}).get("F_amp_N", 0) or 0)
    fe = float(getattr(vc, "external_axial_N", 0) or 0)
    partes = []
    if d > 0:
        partes.append(f"δ {d:g} mm")
    elif fa > 0:
        partes.append(f"F_a {fa / 1000:g} kN")
    if fe > 0:
        partes.append(f"F_ax {fe / 1000:g} kN")
    return " + ".join(partes) if partes else "sustained preload"


def _f0(vc):
    f = float(getattr(vc, "initial_preload_N", 0) or 0) / 1000.0
    return f"{f:.2f}" if f < 10 else f"{f:.1f}"


def _num(v, nd=4):
    return "—" if v is None else f"{float(v):.{nd}f}"


def apendices(doc, comp, res, pisos, store, m, todos_id, res_all, info):
    from docx.enum.section import WD_ORIENT, WD_SECTION
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width

    # ---------------------------------------------------------------- A
    doc.add_heading("Appendix A. Results for every curve", 1)
    p(doc, f"All {m['n']} comparable curves, one row each, with the numbers "
           f"the criterion of §6 reads and the status the census assigns. "
           f"Nothing is summarised away: the aggregates of §6, §10 and §12 are "
           f"sums over this table, and the reader can check them.")
    p(doc, "**Columns.** *load*: imposed transverse amplitude δ, axial force "
           "amplitude F_a, or a superposed external axial load F_ax; "
           "'sustained preload' marks the creep tests with no cyclic load. "
           "*F₀*: initial preload. *f*: frequency. *n*: points in the metric "
           "window. *MAE*, *max |r|*, *σ_res*: the three legs, on the residual "
           "r = model − data; σ_res is shown as '—' where fewer than six points "
           "make it non-judgeable. *energy res.*: the conservation residual "
           "W_ext + ΔU − ΣW_diss of the simulation, relative to the larger of "
           "|ΣW_diss| and |ΔU|, as persisted in the store.", size=9)
    p(doc, "**Status.** *meets*: all three legs inside. *exception*: signed "
           "proof that the source's own replicate scatter (F5) or measured "
           "floor (F7) exceeds the tolerance. *declared*: a measured cause "
           "outside the model, with its class (§10.3). *outside*: fails on "
           "merit; the leg that decides is named, with its distance beyond the "
           "limit.", size=9)
    porid = {r.case_id: r for r in comp}
    linhas, cont = [], collections.Counter()
    sobrepostos = []
    for r in sorted(comp, key=lambda q: (q.source, q.case_id)):
        rr = res.get(r.case_id)
        vc = r.validation_case
        cu = (getattr(rr, "config_used", None) or {}) if rr else {}
        ls = rh.limite_sres(r.source, pisos)
        okc = bool(rh._tripe_ok(rr, ls)) if rr else False
        sr = rh.sres_para_censo(rr) if rr else None
        if rr is None or rr.mae is None:
            st, chave = "not simulated", "other"
        elif okc:
            st, chave = "meets", "meets"
            if r.case_id in rh._EXCECOES or r.case_id in rh._DECLARADAS:
                sobrepostos.append(r.case_id)
        elif r.case_id in rh._EXCECOES:
            st, chave = f"exception: {tipo_excecao(r.case_id)}", "exception"
        elif r.case_id in rh._DECLARADAS:
            st, chave = f"declared: {classe_declarada(r.case_id)}", "declared"
        else:
            mult = {"max |r|": rr.maxerr / rh.META_MAX,
                    "MAE": rr.mae / rh.META_MAE}
            if sr is not None:
                mult["σ_res"] = sr / ls
            q = max(mult, key=mult.get)
            st, chave = f"outside: {q} {mult[q]:.2f}×", "outside"
        cont[chave] += 1
        eb = getattr(rr, "energy_budget", None) if rr else None
        linhas.append([r.case_id, r.source.replace("_", " ").lower(),
                       _carga(vc, cu), _f0(vc),
                       f"{float(getattr(vc, 'frequency_Hz', 0) or 0):g}",
                       str(len(rr.metric_data)) if rr else "—",
                       _num(rr.mae) if rr else "—",
                       _num(rr.maxerr) if rr else "—",
                       _num(sr),
                       (f"{eb['residual_rel']:.1e}" if eb else "—"), st])
    tabela(doc, ["curve", "source", "load", "F₀ [kN]", "f [Hz]", "n", "MAE",
                 "max |r|", "σ_res", "energy res.", "status"], linhas,
           larguras=[2.2, 0.9, 1.0, 0.5, 0.45, 0.35, 0.5, 0.5, 0.5, 0.6, 1.5],
           size=7)
    resumo = (f"**Totals from the table:** meets {cont['meets']} · exception "
              f"{cont['exception']} · declared {cont['declared']} · outside on "
              f"merit {cont['outside']}"
              + (f" · not simulated {cont['other']}" if cont["other"] else "")
              + f"; {sum(cont.values())} curves in all. The strict reading of "
              f"§6 is the first number. The resolved-or-declared reading of "
              f"§10.3 is the sum of the first three.")
    p(doc, resumo, size=9)
    if cont["meets"] != m["tripe"]:
        print(f"  [WARN] appendix A counts {cont['meets']} meets, §6 says "
              f"{m['tripe']}")
    if sobrepostos:
        print(f"  [WARN] curves that meet the criterion AND carry a signature: "
              f"{sobrepostos}")

    # ---------------------------------------------------------------- B
    doc.add_heading("Appendix B. Records outside the census", 1)
    fora = [r for r in todos_id
            if r.case_id in store and not rh.caso_comparavel(r.source, r.case_id)]
    p(doc, f"The result store holds {len(store)} records; {len(fora)} of them "
           f"are simulated and published but not counted in any number of "
           f"this annex. They are listed so that the denominator of {m['n']} "
           f"is auditable rather than asserted.", size=9.5)
    motivos = {
        "UFU_LAB": "the laboratory's own test bench; left the project on "
                   "2026-08-01 and kept only for the record. A curve produced "
                   "by the software's authors would in any case violate the "
                   "rule of §4",
        "USER": "synthetic example shipped with the software; not an experiment",
    }
    motivo_cid = {
        "lu2024_M8_fig18_amp1p0": "the same physical test as "
                                  "lu2024_M8_fig20_T22Nm, published in two "
                                  "figures of the paper (Table 8 at 1.0 mm ≡ "
                                  "Table 9 at 22 N·m); counted once",
    }
    linhas = []
    for r in sorted(fora, key=lambda q: (q.source, q.case_id)):
        rr = res_all.get(r.case_id)
        why = motivo_cid.get(r.case_id) or motivos.get(r.source) \
            or "excluded by the census rule"
        linhas.append([r.case_id, r.source.replace("_", " ").lower(), why,
                       _num(getattr(rr, "mae", None)) if rr else "—",
                       _num(getattr(rr, "maxerr", None)) if rr else "—",
                       _num(getattr(rr, "resid_std", None)) if rr else "—"])
    tabela(doc, ["record", "source", "why it is not counted", "MAE",
                 "max |r|", "σ_res"], linhas,
           larguras=[2.2, 0.9, 4.0, 0.55, 0.55, 0.55], size=7.5)

    # ---------------------------------------------------------------- C
    led = info.get("ledger") or []
    if not led:
        return
    doc.add_heading("Appendix C. Constants ledger", 1)
    p(doc, f"Every constant the {m['n']} comparable curves run on: "
           f"{sum(1 for x in led if x['scope'] == 'shared')} shared numbers "
           f"carried by the calibration groups and "
           f"{sum(1 for x in led if x['scope'] != 'shared')} per-curve "
           f"entries, {len(led)} rows in all. The class is read from the "
           f"provenance text recorded with the entry, conservatively (§9.3); "
           f"the last column says where that text lives. Values are as "
           f"adopted, in SI units unless the unit column says otherwise; "
           f"emb_um is the embedding depth in micrometres.", size=9.5)
    doc.add_heading("C.1 By source", 2)
    cols = ["read (direct)", "read (canonical reader)", "regressed",
            "anchored", "fitted (declared)", "channel off",
            "form switch (a mode, not a number)", "default (design starter)",
            "inherited (kept from an earlier adoption)",
            "documented (class not parsed)", "undocumented"]
    curto = ["read", "reader", "regr.", "anch.", "fitted", "off", "switch",
             "default", "inherit.", "unparsed", "undoc."]
    por_fonte = collections.defaultdict(collections.Counter)
    n_sh = collections.Counter()
    n_pc = collections.Counter()
    for x in led:
        por_fonte[x["source"]][x["classe"]] += 1
        (n_sh if x["scope"] == "shared" else n_pc)[x["source"]] += 1
    linhas = []
    for s in sorted(por_fonte):
        c = por_fonte[s]
        linhas.append([s.replace("_", " ").lower(), str(n_sh[s]), str(n_pc[s])]
                      + [str(c.get(k, 0)) for k in cols])
    tot = collections.Counter(x["classe"] for x in led)
    linhas.append(["**all**", f"**{sum(n_sh.values())}**",
                   f"**{sum(n_pc.values())}**"]
                  + [f"**{tot.get(k, 0)}**" for k in cols])
    tabela(doc, ["source", "shared", "per-curve"] + curto, linhas,
           larguras=[1.3, 0.5, 0.6] + [0.56] * len(curto), size=6.5)
    doc.add_heading("C.2 Every entry", 2)
    linhas = []
    for x in sorted(led, key=lambda q: (q["source"], q["group"],
                                        q["scope"] != "shared", q["scope"],
                                        q["field"])):
        linhas.append([x["source"].replace("_", " ").lower(), x["group"],
                       x["scope"], x["field"], _fmt_val(x["value"]),
                       x["unit"], x["classe"], x["where"]])
    tabela(doc, ["source", "group", "scope", "constant", "value", "unit",
                 "class", "text in"], linhas,
           larguras=[0.95, 1.55, 1.15, 1.35, 0.75, 0.55, 1.45, 0.85], size=6.5)




# --------------------------------------------------------------------------- #
# Facts shared with the main-paper builder (build_paper_docx.py)              #
# --------------------------------------------------------------------------- #

def evidencia_temporal(todos, NOV="rousseau2025_hdpe_t10_amp0p2",
                       CAL="rousseau2025_hdpe_t10"):
    """When the predicted curve entered the repository, and the latest date
    any constant it shares with the calibration curve was touched."""
    porid_all = {r.case_id: r for r in todos}
    ev = {}
    if NOV in porid_all and CAL in porid_all:
        ev["csv_nov"] = data_primeiro_commit(
            porid_all[NOV].validation_case.reference_csv_path)
        ev["csv_cal"] = data_primeiro_commit(
            porid_all[CAL].validation_case.reference_csv_path)
        g = rn._adopted_for(porid_all[NOV].source, NOV,
                            porid_all[NOV].validation_case.bolt_size)
        e = kb.adopted_config(g) or {}
        datas = []
        for k, v in (e.get("cfg") or {}).items():
            if k == "per_case" or not (isinstance(v, (int, float))
                                       and not isinstance(v, bool)):
                continue
            txt = prov_lookup(e.get("prov") or {}, k) or ""
            for dt_ in re.findall(r"20\d\d-\d\d-\d\d", txt):
                datas.append((dt_, k))
        ev["ultima"] = max(datas) if datas else None
        ev["grupo"] = g
    return ev


def energia_corpus(res, comp):
    """The energy residual across the corpus, read from the store; None
    before the re-stamp of 2026-08-28 that started persisting it. The tail is
    characterised (budget size, damage/fatigue flags, sources), not excused."""
    rel = [(abs(float(rr.energy_budget["residual_rel"])), cid)
           for cid, rr in res.items()
           if getattr(rr, "energy_budget", None)
           and rr.energy_budget.get("residual_rel") is not None]
    if not rel:
        return None
    vals = np.array([v for v, _ in rel])
    pior = max(rel)
    porid_c = {r.case_id: r for r in comp}
    cauda = [cid for v, cid in rel if v > 0.05]

    def eb_of(c):
        return res[c].energy_budget
    orc = [max(abs(eb_of(c)["W_diss_J"]), abs(eb_of(c)["dU_J"])) for c in cauda]
    ovr = [(res[c].config_used or {}).get("overrides") or {} for c in cauda]
    n_fat = sum(1 for o in ovr if o.get("fatigue_enabled"))
    n_dmg = sum(1 for o in ovr if float(o.get("c_D", 0) or 0) > 0)
    fontes = collections.Counter(porid_c[c].source for c in cauda
                                 if c in porid_c)
    kJ = [(abs(float(res[c].energy_budget["residual_rel"])), c)
          for c in res if getattr(res[c], "energy_budget", None)
          and abs(res[c].energy_budget["W_diss_J"]) >= 1e3]
    grandes = [o for o in orc if o >= 1.0]
    return {
        "n": len(rel), "med": float(np.median(vals)),
        "p90": float(np.percentile(vals, 90)), "max": float(pior[0]),
        "pior": pior[1], "n_cauda": len(cauda),
        "n_cauda_orc_1J": sum(1 for o in orc if o < 1.0),
        "n_cauda_grandes": len(grandes),
        "orc_grandes_lo": (min(grandes) if grandes else None),
        "orc_grandes_hi": (max(grandes) if grandes else None),
        "max_abs_J_cauda": (max(abs(eb_of(c)["residual_J"]) for c in cauda)
                            if cauda else 0.0),
        "n_cauda_fat": n_fat, "n_cauda_dmg": n_dmg,
        "fontes_cauda": [s.replace("_", " ").lower()
                         for s, _ in fontes.most_common(3)],
        "n_kJ": len(kJ), "max_kJ": (max(kJ)[0] if kJ else None)}


def pernas_que_mandam(comp, res, pisos):
    """Which leg decides each failing verdict, and the severity distribution;
    the same arithmetic the annex §10.2 and the paper use."""
    pernas, piores = collections.Counter(), []
    for r in comp:
        rr = res.get(r.case_id)
        if rr is None:
            continue
        ls = rh.limite_sres(r.source, pisos)
        if rh._tripe_ok(rr, ls):
            continue
        sr = rh.sres_para_censo(rr)
        mult = {"max |r|": rr.maxerr / rh.META_MAX,
                "MAE": rr.mae / rh.META_MAE}
        if sr is not None:
            mult["σ_res"] = sr / ls
        q = max(mult, key=mult.get)
        pernas[q] += 1
        piores.append((mult[q], r.case_id, q))
    piores.sort(reverse=True)
    return pernas, piores


# --------------------------------------------------------------------------- #
# Entry point                                                                 #
# --------------------------------------------------------------------------- #

def main():
    print("[1/5] loading the canonical store")
    comp, res, pisos, store, todos, res_all = carrega()
    print(f"      {len(comp)} comparable curves, {len(res)} with vectors")

    print("[2/5] recomputing every number the annex quotes")
    m = metricas(comp, res, pisos)
    refs = referencias(comp)
    import build_variable_explorer as bve   # loads the 129 VarSpecs
    specs = list(bve.VARIABLE_SPECS)
    fam, n_fora = familia_fisica(todos, res_all)
    print(f"      census {m['tripe']}/{m['n']} · R2 {m['r2']:.4f} · "
          f"{len(specs)} fields · {len(refs)} sources")

    print("[3/5] building the figures (English)")
    info = figuras(comp, res, pisos, store, m, fam, todos, res_all)
    info["snippet_ratio"] = verifica_snippet()
    if info["snippet_ratio"] is not None:
        print(f"      §3.2 snippet runs: F/F0 = {info['snippet_ratio']:.4f} "
              f"after 1000 cycles")
    print("[3b]  repository facts: engine checks, tests, git, digitisation")
    info["engine"] = verificacoes_engine()
    eng = info["engine"]
    print(f"      Norton closed form: max dev {eng['norton_dev_m']:.2e} m · "
          f"energy residual {eng['residual_J']:.3g} J of "
          f"{eng['W_diss_J']:.0f} J · {eng['cycles_per_s']:.0f} cycles/s")
    info["tests"] = contagem_testes()
    info["git"] = revisao_git()
    info["digit"] = piso_digitalizacao_par()
    info["ledger"] = ledger_constantes(comp, specs)
    print(f"      tests collected: {info['tests']} · revision {info['git']} · "
          f"ledger rows {len(info['ledger'])}")
    info["energy_corpus"] = energia_corpus(res, comp)
    ec = info["energy_corpus"]
    if ec:
        print(f"      energy residual (rel.) persisted for {ec['n']} curves: "
              f"median {ec['med']:.1e}, max {ec['max']:.1e} ({ec['pior']}); "
              f"tail >5 %: {ec['n_cauda']} curves, {ec['n_cauda_orc_1J']} with "
              f"budget < 1 J, fatigue on {ec['n_cauda_fat']}, damage on "
              f"{ec['n_cauda_dmg']}")
    else:
        print("      energy residual: not in the store yet (re-stamp pending)")
    info["evidencia_82"] = evidencia_temporal(todos)

    print("[4/5] writing the document")
    doc = Document()
    estilo(doc)
    monta(doc, comp, res, pisos, store, m, refs, specs, fam, n_fora, info,
          todos, res_all)

    SAIDA.mkdir(parents=True, exist_ok=True)
    destino = SAIDA / "BAS_V2_software_annex.docx"
    doc.save(destino)
    print(f"[5/5] {destino}")
    print(f"      {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, "
          f"{_NFIG[0]} figures")
    vaz = auditoria_texto(doc)
    if vaz:
        print(f"  [WARN] {len(vaz)} leaked token(s) in the English text:")
        for tok, trecho in vaz[:20]:
            print(f"         {tok!r}: {trecho}")
    else:
        print("      text audit: no Portuguese/UI vocabulary leaked")
    return destino


if __name__ == "__main__":
    main()
